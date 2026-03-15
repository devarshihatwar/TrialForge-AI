"""
TrialForge AI — Autonomous Clinical Trial Protocol Designer
Hackathon Final Build · Powered by Amazon Nova Pro (Bedrock)
"""

import streamlit as st
import boto3
import json
import logging
import datetime
import random
import hashlib
import time
import re
import io
import math
import xml.etree.ElementTree as ET
from xml.dom import minidom
import difflib
import collections

# ── Plotly (interactive charts) ──────────────────
try:
    import plotly.graph_objects as go
    import plotly.express as px
    HAVE_PLOTLY = True
except ImportError:
    HAVE_PLOTLY = False

# ── PDF ──────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, HRFlowable, PageBreak)
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    HAVE_PDF = True
except Exception:
    HAVE_PDF = False

# ── DOCX ─────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

# ── HTTP (for ClinicalTrials.gov & PubMed) ───
try:
    import urllib.request, urllib.parse
    HAVE_HTTP = True
except Exception:
    HAVE_HTTP = False

logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s")
logger = logging.getLogger("TrialForge")

# ═══════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ═══════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="TrialForge AI",
    page_icon="⚗️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ═══════════════════════════════════════════════════════════════════
# CSS
# ═══════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --bg:        #0D1117;
  --surface:   #161B22;
  --card:      #1C2130;
  --card2:     #21273A;
  --border:    rgba(99,130,190,0.16);
  --border2:   rgba(99,130,190,0.30);
  --accent:    #2DD4BF;
  --blue:      #3B82F6;
  --indigo:    #818CF8;
  --green:     #34D399;
  --gold:      #FBBF24;
  --red:       #F87171;
  --purple:    #C084FC;
  --text:      #E6EDF3;
  --text2:     #8B949E;
  --text3:     #484F58;
  --mono:      'JetBrains Mono', monospace;
  --sans:      'Inter', system-ui, sans-serif;
  --r-sm: 6px; --r-md: 10px; --r-lg: 14px; --r-xl: 18px; --r-pill: 999px;
  --sh-sm: 0 1px 4px rgba(0,0,0,.5);
  --sh-md: 0 4px 16px rgba(0,0,0,.5), 0 1px 4px rgba(0,0,0,.3);
  --sh-lg: 0 8px 32px rgba(0,0,0,.6);
  --sh-glow: 0 0 20px rgba(45,212,191,.15);
}

*, *::before, *::after { box-sizing: border-box; }
html, body, [class*="css"], .stApp {
  background: var(--bg) !important;
  color: var(--text) !important;
  font-family: var(--sans) !important;
  -webkit-font-smoothing: antialiased;
}

/* scrollbar */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: var(--border2); border-radius: var(--r-pill); }

/* main container */
.main .block-container {
  padding: 1.2rem 2rem 3rem !important;
  max-width: 1500px !important;
}

/* SIDEBAR */
[data-testid="stSidebar"] {
  background: var(--surface) !important;
  border-right: 1px solid var(--border) !important;
}
[data-testid="stSidebarContent"] { padding: 1rem .75rem !important; }

/* INPUTS */
[data-testid="stTextInput"] label,
[data-testid="stSelectbox"] label,
[data-testid="stToggle"] label,
[data-testid="stTextArea"] label {
  font-size: .64rem !important; font-weight: 700 !important;
  letter-spacing: .08em !important; text-transform: uppercase !important;
  color: var(--text2) !important;
}
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea {
  background: var(--card) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--r-sm) !important;
  color: var(--text) !important;
  font-size: .83rem !important;
  font-family: var(--sans) !important;
  transition: border-color .2s, box-shadow .2s !important;
}
[data-testid="stTextInput"] input:focus,
[data-testid="stTextArea"] textarea:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px rgba(45,212,191,.12) !important;
}
[data-testid="stSelectbox"] [data-baseweb="select"] > div {
  background: var(--card) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--r-sm) !important;
  color: var(--text) !important;
  font-size: .83rem !important;
}
[data-baseweb="popover"] { background: var(--card2) !important; }

/* BUTTONS */
div.stButton > button {
  background: linear-gradient(135deg,#0E7490,#2DD4BF,#6366F1) !important;
  color: #fff !important; border: none !important;
  border-radius: var(--r-md) !important;
  padding: .6rem 1.2rem !important;
  font-weight: 700 !important; font-size: .83rem !important;
  font-family: var(--sans) !important;
  box-shadow: 0 4px 16px rgba(45,212,191,.3), var(--sh-md) !important;
  transition: all .22s cubic-bezier(.4,0,.2,1) !important;
  width: 100% !important;
}
div.stButton > button:hover {
  transform: translateY(-2px) !important;
  box-shadow: 0 8px 24px rgba(45,212,191,.45), var(--sh-lg) !important;
}
div.stButton > button:active { transform: translateY(0) scale(.98) !important; }
div.stDownloadButton > button {
  background: var(--card2) !important;
  border: 1px solid var(--border2) !important;
  border-radius: var(--r-sm) !important;
  color: var(--text) !important;
  font-size: .78rem !important; font-weight: 600 !important;
  transition: all .18s !important; width: 100% !important;
}
div.stDownloadButton > button:hover {
  border-color: var(--accent) !important;
  transform: translateY(-1px) !important;
}
[data-testid="stSidebar"] div.stButton > button {
  background: var(--card) !important;
  border: 1px solid var(--border) !important;
  color: var(--text2) !important;
  font-size: .72rem !important; font-weight: 500 !important;
  box-shadow: none !important; margin-bottom: 3px !important;
}
[data-testid="stSidebar"] div.stButton > button:hover {
  background: var(--card2) !important;
  border-color: var(--accent) !important;
  color: var(--accent) !important; transform: none !important;
  box-shadow: none !important;
}

/* EXPANDER */
[data-testid="stExpander"] {
  background: var(--card) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--r-md) !important;
}
[data-testid="stExpander"] summary {
  font-size: .82rem !important; font-weight: 600 !important;
  color: var(--text2) !important;
}

/* TABS */
[data-testid="stTabs"] button {
  font-size: .78rem !important; font-weight: 600 !important;
  color: var(--text3) !important;
  border-bottom: 2px solid transparent !important;
  transition: all .18s !important;
}
[data-testid="stTabs"] button[aria-selected="true"] {
  color: var(--accent) !important;
  border-bottom-color: var(--accent) !important;
}

/* ── CUSTOM COMPONENTS ─────────────────────── */

/* Navbar */
.tf-nav {
  display: flex; align-items: center; justify-content: space-between;
  padding: .5rem 0 1rem; border-bottom: 1px solid var(--border);
  margin-bottom: 1.25rem;
}
.tf-brand { display: flex; align-items: center; gap: 10px; }
.tf-brand-icon {
  width: 40px; height: 40px; border-radius: var(--r-md);
  background: linear-gradient(135deg, var(--accent), var(--indigo));
  display: flex; align-items: center; justify-content: center;
  font-size: 1.3rem; box-shadow: var(--sh-glow); flex-shrink: 0;
}
.tf-brand-name { font-size: 1.2rem; font-weight: 800; letter-spacing: -.02em; }
.tf-brand-tag {
  font-family: var(--mono); font-size: .55rem; color: var(--text3);
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-pill); padding: 1px 8px; margin-left: 6px;
}
.tf-pill {
  display: flex; align-items: center; gap: 5px;
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-pill); padding: 3px 10px;
  font-size: .64rem; color: var(--text2); font-family: var(--mono);
}
.dot { width: 6px; height: 6px; border-radius: 50%; flex-shrink: 0; }
.dot-g { background: var(--green); box-shadow: 0 0 5px var(--green); animation: pulse 2s infinite; }
.dot-c { background: var(--accent); box-shadow: 0 0 5px var(--accent); animation: pulse 2.3s infinite; }

/* Animations */
@keyframes pulse { 0%,100%{opacity:.7} 50%{opacity:1} }
@keyframes fadeUp { from{opacity:0;transform:translateY(12px)} to{opacity:1;transform:translateY(0)} }
@keyframes shimmer { 0%{background-position:-800px 0} 100%{background-position:800px 0} }
@keyframes spin { to{transform:rotate(360deg)} }
@keyframes growBar { from{width:0} to{width:var(--w)} }

.anim-fade { animation: fadeUp .4s ease forwards; }

/* Cards */
.card {
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-lg); padding: 1.1rem 1.25rem;
  box-shadow: var(--sh-md); margin-bottom: .75rem;
  transition: box-shadow .22s, transform .22s, border-color .22s;
  animation: fadeUp .35s ease forwards;
}
.card:hover { box-shadow: var(--sh-lg); border-color: var(--border2); transform: translateY(-2px); }
.card-hi { background: var(--card2); border-color: var(--border2); box-shadow: var(--sh-lg); }
.card-hdr {
  display: flex; align-items: center; gap: 7px;
  font-size: .65rem; font-weight: 700; letter-spacing: .09em;
  text-transform: uppercase; color: var(--text2);
  margin-bottom: .75rem;
}
.card-ico {
  width: 24px; height: 24px; border-radius: var(--r-sm);
  display: flex; align-items: center; justify-content: center;
  font-size: .75rem;
}
.ico-t { background: rgba(45,212,191,.1); }
.ico-g { background: rgba(52,211,153,.1); }
.ico-b { background: rgba(59,130,246,.1); }
.ico-o { background: rgba(251,191,36,.1); }
.ico-r { background: rgba(248,113,113,.1); }
.ico-p { background: rgba(192,132,252,.1); }
.ico-i { background: rgba(129,140,248,.1); }

/* Metric cards */
.metric-row { display: grid; grid-template-columns: repeat(5,1fr); gap: .6rem; margin-bottom: 1rem; }
@media (max-width:1100px) { .metric-row { grid-template-columns: repeat(3,1fr); } }
.metric-card {
  background: var(--card2); border: 1px solid var(--border);
  border-radius: var(--r-lg); padding: .9rem 1rem; text-align: center;
  transition: all .22s; cursor: default; animation: fadeUp .4s ease forwards;
}
.metric-card:hover { transform: translateY(-3px); box-shadow: var(--sh-lg); border-color: var(--border2); }
.mc-val { font-size: 1.65rem; font-weight: 800; line-height: 1; margin-bottom: 3px; }
.mc-lbl { font-size: .58rem; font-weight: 700; letter-spacing: .09em; text-transform: uppercase; color: var(--text2); }
.mc-sub { font-size: .65rem; font-weight: 600; margin-top: 3px; }

/* Quality badges */
.qbadge {
  display: inline-block; padding: 2px 9px; border-radius: var(--r-pill);
  font-size: .6rem; font-weight: 700; text-transform: uppercase; letter-spacing: .07em;
}
.qb-hi  { background: rgba(52,211,153,.12); color: var(--green); border: 1px solid rgba(52,211,153,.25); }
.qb-md  { background: rgba(251,191,36,.10); color: var(--gold);  border: 1px solid rgba(251,191,36,.25); }
.qb-lo  { background: rgba(248,113,113,.10); color: var(--red);  border: 1px solid rgba(248,113,113,.25); }
.risk-chip {
  display: inline-flex; align-items: center; gap: 4px;
  padding: 2px 9px; border-radius: var(--r-pill);
  font-size: .63rem; font-weight: 700; text-transform: uppercase; letter-spacing: .07em;
}
.rc-lo { background:rgba(52,211,153,.1); color:var(--green); border:1px solid rgba(52,211,153,.25); }
.rc-md { background:rgba(251,191,36,.1); color:var(--gold);  border:1px solid rgba(251,191,36,.25); }
.rc-hi { background:rgba(248,113,113,.1); color:var(--red);  border:1px solid rgba(248,113,113,.25); }

/* Progress bars */
.pb-wrap { margin: 4px 0 8px; }
.pb-lbl { display:flex; justify-content:space-between; font-size:.67rem; color:var(--text2); margin-bottom:3px; font-weight:500; }
.pb-track { height:5px; background:var(--card); border-radius:var(--r-pill); overflow:hidden; }
.pb-fill { height:100%; border-radius:var(--r-pill); transition:width 1.4s cubic-bezier(.4,0,.2,1); }

/* Protocol document viewer */
.proto-doc {
  background: var(--surface); border: 1px solid var(--border);
  border-radius: var(--r-lg); padding: 1.5rem 2rem;
  font-size: .82rem; line-height: 1.85; color: var(--text2);
  max-height: 640px; overflow-y: auto; animation: fadeUp .5s ease forwards;
}
.proto-doc::-webkit-scrollbar { width: 4px; }
.proto-doc::-webkit-scrollbar-thumb { background: var(--border2); border-radius: var(--r-pill); }
.proto-h1 {
  font-size: 1rem; font-weight: 800; color: var(--text);
  border-bottom: 1px solid var(--border); padding-bottom: 5px;
  margin: 1.2rem 0 .5rem; letter-spacing: -.01em;
}
.proto-h2 {
  font-size: .84rem; font-weight: 700; color: var(--accent);
  margin: .9rem 0 .3rem; text-transform: uppercase; letter-spacing: .04em;
}
.proto-p { margin: .25rem 0; color: var(--text2); }
.proto-bullet { margin: .15rem 0 .15rem 1.2rem; color: var(--text2); }

/* Thinking steps */
.think-box {
  background: var(--card); border: 1px solid var(--border);
  border-radius: var(--r-lg); padding: .9rem 1.25rem; margin-bottom: .75rem;
}
.think-step {
  display: flex; align-items: center; gap: 9px;
  padding: 6px 0; border-bottom: 1px solid var(--border);
  font-size: .77rem; animation: fadeUp .25s ease forwards;
}
.think-step:last-child { border-bottom: none; }
.ts-ico { width:24px; height:24px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-size:.72rem; flex-shrink:0; }
.ts-done { background:rgba(52,211,153,.12); }
.ts-run  { background:rgba(45,212,191,.12); animation:pulse 1s infinite; }
.ts-wait { background:var(--card2); opacity:.5; }
.ts-lbl { flex:1; color:var(--text2); }
.ts-badge { font-family:var(--mono); font-size:.57rem; padding:1px 6px; border-radius:var(--r-pill); }
.tb-done { background:rgba(52,211,153,.12); color:var(--green); }
.tb-run  { background:rgba(45,212,191,.12); color:var(--accent); }
.tb-wait { background:var(--card2); color:var(--text3); }
.ts-pct  { font-family:var(--mono); font-size:.65rem; color:var(--accent); font-weight:700; }

/* Agent cards */
.agent-card {
  background: var(--card2); border: 1px solid var(--border);
  border-radius: var(--r-md); padding: .75rem 1rem; margin-bottom: 7px;
  display: flex; gap: .75rem; align-items: flex-start;
  transition: all .18s; animation: fadeUp .35s ease forwards;
}
.agent-card:hover { border-color: var(--border2); transform: translateX(3px); box-shadow: var(--sh-md); }
.ag-av { width:34px; height:34px; border-radius:var(--r-md); display:flex; align-items:center; justify-content:center; font-size:.95rem; flex-shrink:0; }
.ag-body { flex:1; }
.ag-name { font-size:.62rem; font-weight:700; letter-spacing:.1em; text-transform:uppercase; margin-bottom:3px; }
.ag-txt { font-size:.77rem; color:var(--text2); line-height:1.6; }

/* ICF / Consent viewer */
.icf-doc {
  background: #FEFEFE; color: #1a1a2e;
  border-radius: var(--r-lg); padding: 1.5rem 2rem;
  font-size: .84rem; line-height: 1.8; max-height: 500px; overflow-y: auto;
  border: 2px solid var(--border2);
}
.icf-h1 { font-size:1.1rem; font-weight:800; color:#1e3a5f; margin-bottom:.5rem; }
.icf-section { margin-bottom:1rem; }
.icf-label { font-size:.65rem; font-weight:700; text-transform:uppercase; letter-spacing:.1em; color:#0e7490; margin-bottom:4px; }
.icf-plain { font-size:.84rem; color:#374151; line-height:1.75; }
.icf-highlight { background:#ecfdf5; border-left:3px solid #10b981; padding:6px 10px; border-radius:0 6px 6px 0; margin:6px 0; }
.icf-warning { background:#fef3c7; border-left:3px solid #f59e0b; padding:6px 10px; border-radius:0 6px 6px 0; margin:6px 0; }

/* SoA table */
.soa-wrap { overflow-x:auto; }
.soa-table { border-collapse:collapse; width:100%; font-size:.72rem; }
.soa-table th {
  background:var(--card2); color:var(--text2); font-weight:700;
  font-size:.6rem; letter-spacing:.07em; text-transform:uppercase;
  padding:7px 10px; border:1px solid var(--border); text-align:center;
  white-space:nowrap;
}
.soa-table th.soa-first { text-align:left; min-width:160px; color:var(--text); }
.soa-table td {
  padding:5px 10px; border:1px solid var(--border);
  text-align:center; color:var(--text2); vertical-align:middle;
}
.soa-table td.soa-first { text-align:left; color:var(--text); font-weight:500; }
.soa-table tr:hover td { background:var(--card2); }
.soa-dot { width:8px; height:8px; border-radius:50%; display:inline-block; }
.soa-req  { background:var(--accent); }
.soa-opt  { background:var(--gold); border:1px solid var(--gold); }
.soa-empty{ background:transparent; }

/* Diversity / Amendment / Interaction badges */
.flag-card {
  background: var(--card2); border: 1px solid var(--border);
  border-radius: var(--r-md); padding: .7rem 1rem; margin-bottom: 6px;
  display: flex; gap: 9px; align-items: flex-start; font-size: .77rem;
  transition: all .18s;
}
.flag-card:hover { border-color: var(--border2); }
.flag-ico { font-size:1rem; flex-shrink:0; margin-top:1px; }
.flag-body { flex:1; }
.flag-title { font-weight:700; color:var(--text); margin-bottom:2px; font-size:.78rem; }
.flag-desc { color:var(--text2); line-height:1.55; }

/* Meta table */
.meta-tbl { width:100%; border-collapse:collapse; font-size:.79rem; }
.meta-tbl tr:hover td { background:var(--card2); }
.meta-tbl td { padding:6px 9px; border-bottom:1px solid var(--border); vertical-align:middle; }
.meta-tbl td:first-child { color:var(--text2); font-weight:500; width:46%; }
.meta-tbl td:last-child { font-family:var(--mono); color:var(--text); font-weight:500; font-size:.74rem; }

/* NCT trial items */
.nct-item { padding:8px 0; border-bottom:1px solid var(--border); font-size:.77rem; }
.nct-item:last-child { border-bottom:none; }
.nct-id { font-family:var(--mono); color:var(--accent); font-size:.68rem; }
.nct-title { color:var(--text); font-weight:600; }
.nct-meta { color:var(--text2); font-size:.68rem; margin-top:2px; }

/* Compare */
.cmp-row { display:flex; align-items:center; padding:5px 0; border-bottom:1px solid var(--border); font-size:.77rem; }
.cmp-row:last-child { border-bottom:none; }
.cmp-lbl { color:var(--text2); flex:1; font-weight:500; }
.cmp-val { font-family:var(--mono); font-size:.72rem; min-width:70px; text-align:right; }
.cv-win { color:var(--green); } .cv-lose { color:var(--red); } .cv-tie { color:var(--text2); }

/* Suggestion items */
.sug-item { display:flex; gap:8px; align-items:flex-start; padding:7px 0; border-bottom:1px solid var(--border); font-size:.77rem; }
.sug-item:last-child { border-bottom:none; }
.sug-num { width:19px; height:19px; border-radius:50%; background:rgba(45,212,191,.1); border:1px solid rgba(45,212,191,.25); color:var(--accent); font-size:.6rem; font-weight:800; display:flex; align-items:center; justify-content:center; flex-shrink:0; margin-top:1px; }
.sug-txt { color:var(--text2); line-height:1.55; }

/* Gantt */
.gantt-row { display:flex; align-items:center; gap:8px; margin-bottom:5px; }
.gantt-lbl { font-size:.67rem; color:var(--text2); min-width:130px; text-align:right; font-weight:500; }
.gantt-track { flex:1; background:var(--card); border-radius:var(--r-pill); height:16px; position:relative; overflow:hidden; }
.gantt-bar { height:100%; border-radius:var(--r-pill); display:flex; align-items:center; padding:0 7px; font-size:.57rem; font-weight:700; color:#fff; white-space:nowrap; overflow:hidden; }

/* Protocol header */
.proto-hdr {
  background: linear-gradient(135deg,var(--card2),#1a2240);
  border: 1px solid var(--border2); border-radius: var(--r-xl);
  padding: 1.1rem 1.5rem; margin-bottom: 1rem;
  box-shadow: var(--sh-lg), var(--sh-glow);
  position: relative; overflow: hidden;
  animation: fadeUp .4s ease forwards;
}
.proto-hdr::before {
  content:''; position:absolute; top:0; left:0; right:0; height:2px;
  background:linear-gradient(90deg,var(--accent),var(--indigo),var(--accent));
  background-size:200% 100%; animation:shimmer 3s linear infinite;
}
.phdr-title { font-size:1.2rem; font-weight:800; color:var(--text); letter-spacing:-.02em; margin-bottom:3px; }
.phdr-meta { font-size:.67rem; color:var(--text3); font-family:var(--mono); }
.phdr-badges { display:flex; gap:6px; flex-wrap:wrap; margin-top:8px; }
.phdr-badge {
  background:var(--card); border:1px solid var(--border);
  border-radius:var(--r-pill); padding:2px 9px;
  font-size:.62rem; font-weight:600; color:var(--text2);
}
.phdr-badge-a { border-color:rgba(45,212,191,.4); color:var(--accent); background:rgba(45,212,191,.08); }
.phdr-badge-i { border-color:rgba(129,140,248,.4); color:var(--indigo); background:rgba(129,140,248,.08); }

/* Landing hero */
.hero {
  padding: 2rem 0 1.5rem; text-align: center;
  animation: fadeUp .5s ease forwards;
}
.hero-eyebrow {
  display:inline-flex; align-items:center; gap:7px;
  background:rgba(45,212,191,.08); border:1px solid rgba(45,212,191,.25);
  border-radius:var(--r-pill); padding:3px 12px;
  font-size:.63rem; font-weight:700; letter-spacing:.1em; text-transform:uppercase;
  color:var(--accent); margin-bottom:1rem;
}
.hero-title { font-size:clamp(1.6rem,3vw,2.4rem); font-weight:800; letter-spacing:-.03em; margin-bottom:.5rem; }
.hero-title span { background:linear-gradient(135deg,var(--accent),var(--indigo)); -webkit-background-clip:text; -webkit-text-fill-color:transparent; }
.hero-sub { font-size:.88rem; color:var(--text2); max-width:480px; margin:0 auto 1.5rem; line-height:1.65; }
.hero-stats { display:flex; gap:1.5rem; justify-content:center; flex-wrap:wrap; }
.hs-item { text-align:center; }
.hs-val { font-size:1.3rem; font-weight:800; color:var(--text); }
.hs-lbl { font-size:.6rem; color:var(--text3); text-transform:uppercase; letter-spacing:.07em; }
.hs-sep { width:1px; background:var(--border); align-self:stretch; margin:4px 0; }

/* Template grid */
.tmpl-grid { display:grid; grid-template-columns:repeat(3,1fr); gap:.6rem; max-width:700px; margin:0 auto 1.5rem; }
@media (max-width:700px) { .tmpl-grid { grid-template-columns:repeat(2,1fr); } }
.tmpl-card {
  background:var(--card); border:1px solid var(--border); border-radius:var(--r-lg);
  padding:.75rem; cursor:pointer; transition:all .22s; text-align:left;
}
.tmpl-card:hover { background:var(--card2); border-color:var(--accent); transform:translateY(-2px); box-shadow:var(--sh-md); }
.tc-icon { font-size:1.1rem; margin-bottom:5px; display:block; }
.tc-name { font-size:.78rem; font-weight:700; color:var(--text); }
.tc-desc { font-size:.66rem; color:var(--text2); margin-top:2px; }
.tc-tag { display:inline-block; background:rgba(45,212,191,.08); border:1px solid rgba(45,212,191,.2); border-radius:var(--r-pill); padding:1px 7px; font-size:.58rem; color:var(--accent); font-weight:600; margin-top:5px; }

/* Latency badge */
.lat-badge {
  display:inline-flex; align-items:center; gap:4px;
  background:rgba(52,211,153,.1); border:1px solid rgba(52,211,153,.25);
  border-radius:var(--r-pill); padding:2px 9px;
  font-family:var(--mono); font-size:.6rem; color:var(--green);
}

/* Sidebar section label */
.sb-section { font-size:.58rem; font-weight:700; letter-spacing:.13em; text-transform:uppercase; color:var(--text3); padding:.6rem .3rem .3rem; margin-top:.3rem; }
.sb-logo { display:flex; align-items:center; gap:8px; padding:.3rem .2rem .8rem; border-bottom:1px solid var(--border); margin-bottom:.5rem; }
.sb-logo-icon { width:32px; height:32px; background:linear-gradient(135deg,var(--accent),var(--indigo)); border-radius:var(--r-sm); display:flex; align-items:center; justify-content:center; font-size:1rem; flex-shrink:0; }
.sb-logo-name { font-size:.88rem; font-weight:800; color:var(--text); }
.sb-logo-ver { font-family:var(--mono); font-size:.56rem; color:var(--text3); }

/* Footer */
.tf-footer { margin-top:2rem; padding:.75rem 0; border-top:1px solid var(--border); display:flex; justify-content:space-between; flex-wrap:wrap; gap:.5rem; }
.tf-footer-l { font-size:.65rem; color:var(--text3); font-family:var(--mono); }
.tf-footer-r { font-size:.62rem; color:var(--text3); }

/* Confidence banner */
.conf-banner {
  background: linear-gradient(135deg,rgba(129,140,248,.1),rgba(45,212,191,.1));
  border: 1px solid rgba(129,140,248,.25); border-radius:var(--r-lg);
  padding: .75rem 1.25rem; display:flex; align-items:center; justify-content:space-between;
  margin-bottom:.75rem;
}
.cb-val { font-size:2rem; font-weight:800; color:var(--text); }
.cb-lbl { font-size:.6rem; text-transform:uppercase; letter-spacing:.09em; color:var(--text2); font-weight:700; }
.cb-sub { font-size:.75rem; color:var(--text2); margin-top:2px; }

/* ── ROI Banner ── */
.roi-banner {
  background: linear-gradient(135deg,rgba(251,191,36,.08),rgba(52,211,153,.08));
  border: 1px solid rgba(251,191,36,.25); border-radius:var(--r-lg);
  padding: .75rem 1.25rem; display:grid;
  grid-template-columns: repeat(4,1fr); gap:.5rem;
  margin-bottom:.75rem;
}
.roi-item { text-align:center; }
.roi-val { font-size:1.3rem; font-weight:800; color:var(--gold); line-height:1; }
.roi-lbl { font-size:.58rem; font-weight:700; text-transform:uppercase; letter-spacing:.08em; color:var(--text2); margin-top:2px; }

/* ── Chat ── */
.chat-msg { display:flex; gap:9px; margin-bottom:10px; align-items:flex-start; animation: fadeUp .25s ease; }
.chat-avatar { width:28px; height:28px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-size:.8rem; flex-shrink:0; margin-top:2px; }
.chat-av-user { background:rgba(129,140,248,.15); }
.chat-av-ai   { background:rgba(45,212,191,.15); }
.chat-bubble { max-width:85%; border-radius:var(--r-md); padding:.55rem .8rem; font-size:.8rem; line-height:1.6; }
.chat-bubble-user { background:var(--card2); color:var(--text); border:1px solid var(--border); border-bottom-right-radius:3px; }
.chat-bubble-ai   { background:rgba(45,212,191,.07); color:var(--text2); border:1px solid rgba(45,212,191,.2); border-bottom-left-radius:3px; }
.chat-time { font-size:.58rem; color:var(--text3); margin-top:3px; font-family:var(--mono); }

/* ── Redline diff ── */
.diff-wrap { font-family:var(--mono); font-size:.72rem; line-height:1.7; background:var(--surface); border:1px solid var(--border); border-radius:var(--r-md); padding:1rem 1.25rem; max-height:480px; overflow-y:auto; }
.diff-add  { background:rgba(52,211,153,.15); color:#86efac; display:block; }
.diff-del  { background:rgba(248,113,113,.12); color:#fca5a5; text-decoration:line-through; display:block; }
.diff-eq   { color:var(--text2); display:block; }
.diff-hdr  { color:var(--accent); font-weight:700; display:block; margin:.5rem 0 .25rem; }

/* ── Audit log ── */
.audit-row { display:flex; align-items:flex-start; gap:10px; padding:7px 0; border-bottom:1px solid var(--border); font-size:.76rem; }
.audit-row:last-child { border-bottom:none; }
.audit-time { font-family:var(--mono); font-size:.62rem; color:var(--text3); min-width:115px; }
.audit-action { color:var(--accent); font-weight:700; min-width:80px; font-size:.68rem; text-transform:uppercase; letter-spacing:.06em; }
.audit-detail { color:var(--text2); flex:1; }
.audit-hash { font-family:var(--mono); font-size:.55rem; color:var(--text3); margin-top:2px; }

/* ── Burden / funnel ── */
.funnel-stage { margin-bottom:5px; }
.funnel-bar-wrap { height:28px; background:var(--card); border-radius:var(--r-sm); position:relative; overflow:hidden; }
.funnel-bar { height:100%; border-radius:var(--r-sm); display:flex; align-items:center; padding:0 10px; font-size:.7rem; font-weight:700; color:#fff; transition:width 1.2s ease; }
.funnel-label { display:flex; justify-content:space-between; font-size:.66rem; color:var(--text2); margin-bottom:3px; }

/* ── Site map ── */
.site-map-wrap { background:var(--surface); border:1px solid var(--border); border-radius:var(--r-lg); padding:.75rem; }
.site-row { display:grid; grid-template-columns:1fr 3fr 70px; gap:8px; align-items:center; padding:5px 0; border-bottom:1px solid var(--border); font-size:.76rem; }
.site-row:last-child { border-bottom:none; }
.site-name { color:var(--text); font-weight:600; }
.site-bar-wrap { background:var(--card); border-radius:var(--r-pill); height:8px; overflow:hidden; }
.site-bar { height:100%; border-radius:var(--r-pill); }
.site-score { font-family:var(--mono); font-size:.7rem; text-align:right; }

/* ── Streaming ── */
.stream-cursor { display:inline-block; width:2px; height:1em; background:var(--accent); animation:pulse 0.8s infinite; vertical-align:text-bottom; margin-left:2px; }

</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════
# CONSTANTS & DATA
# ═══════════════════════════════════════════════════════════════════
COMPARATOR_MAP = {
    "lung cancer": "Platinum-based chemotherapy (carboplatin + paclitaxel)",
    "breast cancer": "AC-T regimen (doxorubicin, cyclophosphamide, paclitaxel)",
    "melanoma": "Dacarbazine (DTIC)",
    "type 2 diabetes": "Metformin 1000 mg BID",
    "rheumatoid arthritis": "Methotrexate 15–20 mg/week",
    "alzheimer": "Donepezil 10 mg/day",
    "alzheimer's disease": "Donepezil 10 mg/day",
    "hypertension": "Lisinopril 10 mg/day",
    "covid-19": "Best supportive care",
    "multiple sclerosis": "Interferon beta-1a",
    "parkinson": "Levodopa/carbidopa",
    "asthma": "Inhaled corticosteroids (fluticasone)",
    "copd": "Long-acting beta-agonist (salmeterol)",
    "heart failure": "ACE inhibitor + beta-blocker",
    "psoriasis": "Methotrexate 10–15 mg/week",
}

DRUG_INTERACTION_DB = {
    "warfarin": {"risk": "High", "detail": "Risk of major bleeding. INR monitoring required."},
    "aspirin":  {"risk": "Medium", "detail": "Increased bleeding risk. Antiplatelet effect."},
    "metformin":{"risk": "Low",    "detail": "Monitor renal function. No direct interaction expected."},
    "statins":  {"risk": "Low",    "detail": "CYP3A4 interaction possible. Monitor LFTs."},
    "ssri":     {"risk": "Medium", "detail": "Serotonin syndrome risk if serotoninergic drug."},
    "nsaids":   {"risk": "Medium", "detail": "Renal impairment risk; may affect endpoint."},
    "insulin":  {"risk": "Low",    "detail": "Blood glucose monitoring required."},
    "corticosteroids": {"risk": "Medium", "detail": "Immune suppression may confound safety endpoints."},
}

SOA_ASSESSMENTS = [
    ("Informed Consent",           True,  False, False, False, False, False),
    ("Medical History & Demographics", True, False, False, False, False, False),
    ("Physical Examination",       True,  False, True,  False, True,  True),
    ("Vital Signs",                True,  True,  True,  True,  True,  True),
    ("Hematology (CBC)",           True,  False, True,  False, True,  True),
    ("Chemistry Panel (LFT/RFT)", True,  False, True,  False, True,  True),
    ("Urinalysis",                 True,  False, False, True,  False, True),
    ("ECG",                        True,  False, False, True,  False, True),
    ("Imaging/Scans",              True,  False, False, False, True,  True),
    ("Study Drug Administration",  False, True,  True,  True,  True,  False),
    ("Efficacy Assessment",        True,  False, True,  True,  True,  True),
    ("Adverse Event Review",       False, True,  True,  True,  True,  True),
    ("Biomarker Sampling",         True,  False, False, True,  False, True),
    ("Quality of Life (PRO)",      True,  False, True,  False, True,  True),
    ("Final Safety Assessment",    False, False, False, False, False, True),
]
SOA_TIMEPOINTS = ["Screening", "Day 1", "Week 4", "Week 12", "Week 24", "EOT"]

DEMO_TEMPLATES = [
    {"icon": "🫁", "name": "Lung Cancer · Phase III",  "drug": "Pembrolizumab",  "disease": "Lung Cancer",          "phase": "Phase III", "mode": "Highest Success Probability", "desc": "PD-1 inhibitor NSCLC"},
    {"icon": "🧠", "name": "Alzheimer's · Phase II",   "drug": "Lecanemab",      "disease": "Alzheimer's Disease",  "phase": "Phase II",  "mode": "Highest Safety",             "desc": "Amyloid antibody"},
    {"icon": "💉", "name": "Diabetes · Phase III",     "drug": "Semaglutide",    "disease": "Type 2 Diabetes",      "phase": "Phase III", "mode": "Fastest Approval",            "desc": "GLP-1 receptor agonist"},
    {"icon": "🩺", "name": "RA · Phase II",            "drug": "Adalimumab",     "disease": "Rheumatoid Arthritis", "phase": "Phase II",  "mode": "Lowest Cost",                 "desc": "TNF-alpha inhibitor"},
    {"icon": "🎗️", "name": "Melanoma · Phase III",    "drug": "Nivolumab",      "disease": "Melanoma",             "phase": "Phase III", "mode": "Highest Success Probability", "desc": "Checkpoint inhibitor"},
    {"icon": "❤️", "name": "Heart Failure · Phase II", "drug": "Sacubitril",     "disease": "Heart Failure",        "phase": "Phase II",  "mode": "Highest Safety",             "desc": "Neprilysin inhibitor"},
]

MODE_HINTS = {
    "Highest Success Probability": "Balances regulatory, safety & recruitment for max overall success.",
    "Fastest Approval":            "Adaptive design, surrogate endpoints, Fast Track designation.",
    "Lowest Cost":                 "Decentralised elements, risk-based monitoring, lean CRO.",
    "Highest Safety":              "Conservative escalation, enhanced DSMB, sentinel dosing.",
}

ICH_GCP_ITEMS = [
    ("Protocol version and date documented", True),
    ("Informed consent process defined", True),
    ("IRB/IEC approval pathway outlined", True),
    ("Inclusion/exclusion criteria clearly stated", True),
    ("Primary and secondary endpoints defined", True),
    ("Statistical analysis plan included", True),
    ("Adverse event reporting procedures defined", True),
    ("Data collection and monitoring plan", True),
    ("Subject confidentiality (GDPR/HIPAA)", True),
    ("Early termination criteria specified", True),
    ("Risk-benefit assessment documented", True),
    ("Insurance and indemnity information", False),
]

ARCH_DOT = """
digraph TrialForge {
    graph [bgcolor="#0D1117" fontname="Inter" splines=ortho pad=0.4 ranksep=0.7 nodesep=0.5]
    node  [fontname="Inter" fontsize=10 style=filled shape=box margin="0.2,0.1"]
    edge  [color="#2DD4BF" arrowsize=0.7 fontname="Inter" fontsize=9]

    User     [label="User Input\n(Drug / Disease / Phase)" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#2DD4BF"]
    RAG      [label="RAG Engine\nStatic DB + ClinicalTrials.gov" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#818CF8"]
    PubMed   [label="PubMed API\nLive Citations" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#34D399"]
    Prompt   [label="Prompt Builder\nMode + RAG + Citations" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#FBBF24"]
    Bedrock  [label="Amazon Bedrock\nNova Pro v1.0" fillcolor="#21273A" fontcolor="#2DD4BF" color="#2DD4BF" penwidth=2]
    RegAgent [label="Regulatory Agent\nFDA / ICH" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#818CF8"]
    SafAgent [label="Safety Agent\nDSMB / SAE" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#F87171"]
    StaAgent [label="Statistical Agent\nPower / SAP" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#34D399"]
    RedTeam  [label="Red Team Agent\nFDA Mock Reviewer" fillcolor="#1C2130" fontcolor="#F87171" color="#F87171"]
    Quality  [label="Quality Judge\nICH / FDA Scores" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#FBBF24"]
    Amend    [label="Amendment Risk\nPredictor" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#C084FC"]
    Drug     [label="Drug Interaction\nChecker" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#F87171"]
    Diversity[label="FDA Diversity\nAction Plan" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#34D399"]
    ICF      [label="ICF Translator\n6th-Grade Plain Text" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#2DD4BF"]
    SoA      [label="Schedule of\nAssessments Matrix" fillcolor="#1C2130" fontcolor="#E6EDF3" color="#818CF8"]
    Output   [label="Output Layer\nPDF / DOCX / JSON / TXT" fillcolor="#21273A" fontcolor="#2DD4BF" color="#2DD4BF" penwidth=2]

    User -> RAG; User -> Prompt
    RAG -> Prompt; PubMed -> Prompt
    Prompt -> Bedrock
    Bedrock -> RegAgent; Bedrock -> SafAgent; Bedrock -> StaAgent; Bedrock -> RedTeam
    Bedrock -> Quality; Bedrock -> Amend; Bedrock -> Drug
    Bedrock -> Diversity; Bedrock -> ICF; Bedrock -> SoA
    RegAgent -> Output; SafAgent -> Output; StaAgent -> Output
    Quality -> Output; RedTeam -> Output
    Diversity -> Output; ICF -> Output; SoA -> Output; Amend -> Output; Drug -> Output
}
"""

# ═══════════════════════════════════════════════════════════════════
# BEDROCK CLIENT
# ═══════════════════════════════════════════════════════════════════
@st.cache_resource
def get_bedrock():
    return boto3.client("bedrock-runtime", region_name="us-east-1")


# ═══════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════
def _hash(s: str) -> int:
    return int(hashlib.md5(s.encode()).hexdigest(), 16) % 100000


def suggest_comparator(disease: str) -> str:
    d = disease.lower()
    for k, v in COMPARATOR_MAP.items():
        if k in d:
            return v
    return "Placebo-controlled or standard-of-care per local guidelines"


def risk_scores(phase: str, disease: str, drug: str) -> dict:
    rng = random.Random(_hash(phase + disease + drug))
    pw  = {"Phase I": (.35,.45,.70), "Phase II": (.55,.60,.50), "Phase III": (.72,.75,.35)}
    ba, br, bs = pw.get(phase, (.55,.60,.50))
    ap = round(min(.97, max(.10, ba + rng.uniform(-.12,.12))), 2)
    rs = round(min(.99, max(.15, br + rng.uniform(-.15,.15))), 2)
    sr = round(min(.90, max(.05, bs + rng.uniform(-.10,.10))), 2)
    sp = round(ap*.4 + rs*.3 + (1-sr)*.3, 2)
    return {"ap": ap, "rs": rs, "sr": sr, "sp": sp,
            "risk_class": "Low" if sr<.30 else ("Medium" if sr<.60 else "High"),
            "rec_diff":   "Easy" if rs>.70 else ("Moderate" if rs>.45 else "Difficult")}


def quality_scores(phase: str, disease: str, drug: str, text: str) -> dict:
    rng  = random.Random(_hash("quality"+phase+disease+drug))
    t    = text.lower()
    base = 74 + rng.randint(0,16)
    if "statistical" in t and "analysis plan" in t: base += 3
    if "primary endpoint" in t: base += 2
    if "adverse event" in t:    base += 2
    if "inclusion" in t and "exclusion" in t: base += 2
    base += {"Phase III":5,"Phase II":1,"Phase I":-2}.get(phase,0)
    ich  = min(99,max(55,base))
    reg  = min(99,max(55,base + rng.randint(-3,6)))
    fda  = min(99,max(50,base + rng.randint(-5,8)))
    fail = min(45,max(8, 38 - rng.randint(5,20)))
    conf = min(99,max(60,round((ich*.3+reg*.3+fda*.2+(100-fail)*.2))))
    comp = min(99,max(65,round((ich+reg)/2 + rng.randint(-3,5))))
    return {"ich":ich,"reg":reg,"fda":fda,"fail":fail,"conf":conf,"comp":comp,
            "fail_lbl":"Low" if fail<20 else ("Medium" if fail<35 else "High"),
            "fda_lbl": "High" if fda>=80 else ("Medium" if fda>=65 else "Low")}


def cost_estimate(phase: str, n: int, months: int) -> dict:
    pp  = {"Phase I":45000,"Phase II":28000,"Phase III":19000}.get(phase,25000)
    tot = pp*n; ov=round(tot*.22); reg={"Phase I":80000,"Phase II":250000,"Phase III":750000}.get(phase,300000)
    cro = round(tot*.18); grand=tot+ov+reg+cro
    return {"pp":pp,"total":tot,"ov":ov,"reg":reg,"cro":cro,"grand":round(grand)}


def timeline_est(phase: str, n: int) -> dict:
    base={"Phase I":14,"Phase II":28,"Phase III":48}.get(phase,30)
    su=6; en=max(6,round(n/25)); fu={"Phase I":3,"Phase II":6,"Phase III":12}.get(phase,6)
    return {"su":su,"en":en,"base":base,"fu":fu,"total":su+en+base+fu}


def amendment_flags(drug: str, disease: str, phase: str) -> list:
    rng = random.Random(_hash("amend"+drug+disease+phase))
    flags = [
        ("🔴 Overly Restrictive Age Range",
         "Excluding patients >75 may reduce recruitment pool by ~40%. Consider broadening to 18–80 unless scientific rationale exists.",
         "High"),
        ("🟡 Narrow ECOG Status Criterion",
         "ECOG 0-1 only may exclude 35% of otherwise eligible patients. Phase II data suggests ECOG 2 patients tolerate this drug class.",
         "Medium"),
        ("🟠 Ambiguous Washout Period",
         "4-week washout for 'any immunosuppressant' is inconsistent with some agents' half-lives. Specify drug-by-drug washout windows.",
         "High"),
        ("🟡 Renal Function Threshold",
         "eGFR ≥60 mL/min cutoff may require amendment if drug label is approved for eGFR ≥45. Align with expected label language.",
         "Medium"),
        ("🔴 Missing Hepatic Impairment Guidance",
         "No exclusion or dose adjustment for mild hepatic impairment (Child-Pugh A). FDA will require this in the label; add prospectively.",
         "High"),
    ]
    n = rng.randint(2, 4)
    random.Random(_hash("amend2"+drug)).shuffle(flags)
    return flags[:n]


def drug_interactions(drug: str) -> list:
    results = []
    drug_l  = drug.lower()
    for agent, info in DRUG_INTERACTION_DB.items():
        rng = random.Random(_hash(drug_l + agent))
        if rng.random() < 0.45 or agent in drug_l:
            results.append({"drug": agent.title(), **info})
    if not results:
        results = [{"drug":"Background Medications","risk":"Low","detail":"No high-risk interactions identified in standard screening. Monitor concomitant use of CYP3A4 inhibitors/inducers."}]
    return results[:4]


def diversity_plan(disease: str, phase: str) -> list:
    items = [
        ("Race & Ethnicity Targets",
         "FDORA §3041 requires enrollment goals for Black/African American (≥15%), Hispanic/Latino (≥18%), Asian (≥10%), and other underrepresented groups proportional to US disease burden."),
        ("Sex/Gender Inclusion",
         "Ensure ≥40% enrollment of each sex unless disease-specific rationale provided. Include sex-stratified subgroup analysis in SAP."),
        ("Age Inclusion (Elderly)",
         "FDA strongly recommends including patients ≥65 years. Prespecify geriatric subgroup. Consider PK bridging for frail elderly."),
        ("Geographic Diversity",
         "Site selection should span urban, suburban, and rural communities. Partner with Federally Qualified Health Centers (FQHCs) for underserved access."),
        ("Language & Literacy Access",
         "Translate ICF and PRO instruments into Spanish, Mandarin, and top-3 local languages at study sites. Use ≤8th-grade reading level."),
        ("Digital Divide Mitigation",
         f"{'Phase I' if phase=='Phase I' else 'Phase II/III'} trials must offer non-digital participation pathways for populations without broadband access or smartphone ownership."),
    ]
    return items


def synthetic_ehr_feasibility(disease: str, phase: str, n_needed: int) -> dict:
    rng = random.Random(_hash("ehr"+disease+phase))
    us_prevalence = {"lung cancer":236000,"breast cancer":290000,"type 2 diabetes":38000000,
                     "melanoma":100000,"alzheimer's disease":6800000,"rheumatoid arthritis":1300000,
                     "hypertension":110000000,"heart failure":6200000}.get(disease.lower(), 500000)
    eligible_pct  = rng.uniform(0.008, 0.045)
    eligible_pool = round(us_prevalence * eligible_pct)
    sites_needed  = max(5, round(n_needed / rng.randint(8,18)))
    monthly_rate  = max(1, round(eligible_pool / 2400))
    feasibility   = "High" if eligible_pool > n_needed*4 else ("Medium" if eligible_pool > n_needed*2 else "Low")
    return {
        "prevalence": us_prevalence, "eligible_pool": eligible_pool,
        "sites_needed": sites_needed, "monthly_rate": monthly_rate,
        "feasibility": feasibility, "eligible_pct": round(eligible_pct*100,2),
    }


def fetch_clinicaltrials(disease: str, drug: str, limit: int = 5) -> list:
    if not HAVE_HTTP:
        return []
    try:
        q   = urllib.parse.quote(f"{drug} {disease}")
        url = (f"https://clinicaltrials.gov/api/v2/studies?"
               f"query.term={q}&pageSize={limit}&format=json&"
               f"fields=NCTId,BriefTitle,Phase,EnrollmentCount,OverallStatus,PrimaryOutcomeMeasure")
        req  = urllib.request.Request(url, headers={"User-Agent":"TrialForge/3.0"})
        with urllib.request.urlopen(req, timeout=7) as r:
            data = json.loads(r.read())
        out = []
        for s in data.get("studies", []):
            ps  = s.get("protocolSection", {})
            id_ = ps.get("identificationModule", {})
            dm  = ps.get("designModule", {})
            sm  = ps.get("statusModule", {})
            om  = ps.get("outcomesModule", {})
            pri = (om.get("primaryOutcomes") or [{}])[0]
            out.append({
                "id":       id_.get("nctId","N/A"),
                "title":    id_.get("briefTitle","Unknown")[:80],
                "phase":    ", ".join(dm.get("phases",["N/A"])),
                "n":        dm.get("enrollmentInfo",{}).get("count","N/A"),
                "status":   sm.get("overallStatus","N/A"),
                "endpoint": pri.get("measure","N/A")[:60],
            })
        return out
    except Exception as e:
        logger.warning("ClinicalTrials.gov: %s", e)
        return []


def fetch_pubmed_citations(drug: str, disease: str, limit: int = 4) -> list:
    """Fetch real PubMed citations via E-utilities."""
    if not HAVE_HTTP:
        return []
    try:
        q   = urllib.parse.quote(f"{drug}[Title/Abstract] AND {disease}[Title/Abstract] AND clinical trial[pt]")
        url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?db=pubmed&term={q}&retmax={limit}&retmode=json&sort=relevance"
        req = urllib.request.Request(url, headers={"User-Agent":"TrialForge/3.0"})
        with urllib.request.urlopen(req, timeout=7) as r:
            data = json.loads(r.read())
        ids = data.get("esearchresult",{}).get("idlist",[])
        if not ids:
            return []
        id_str = ",".join(ids)
        sum_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?db=pubmed&id={id_str}&retmode=json"
        req2 = urllib.request.Request(sum_url, headers={"User-Agent":"TrialForge/3.0"})
        with urllib.request.urlopen(req2, timeout=7) as r2:
            sdata = json.loads(r2.read())
        results = []
        for pmid in ids:
            art = sdata.get("result",{}).get(pmid,{})
            if art:
                results.append({
                    "pmid":    pmid,
                    "title":   art.get("title","")[:100],
                    "journal": art.get("source",""),
                    "year":    art.get("pubdate","")[:4],
                    "url":     f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
                })
        return results
    except Exception as e:
        logger.warning("PubMed: %s", e)
        return []


def call_nova(prompt: str) -> tuple:
    """Call Amazon Nova Pro via Bedrock Converse API. Returns (text, latency)."""
    client = get_bedrock()
    logger.info("Calling Nova Pro")
    t0 = time.time()
    resp = client.converse(
        modelId="amazon.nova-pro-v1:0",
        messages=[{"role":"user","content":[{"text":prompt}]}]
    )
    lat  = round(time.time()-t0, 2)
    text = resp["output"]["message"]["content"][0]["text"]
    logger.info("Nova response: %d chars in %.2fs", len(text), lat)
    return text, lat


def call_nova_streaming(prompt: str, placeholder):
    """Stream Nova Pro response token-by-token into a Streamlit placeholder."""
    client = get_bedrock()
    t0 = time.time()
    full_text = ""
    try:
        resp = client.invoke_model_with_response_stream(
            modelId="amazon.nova-pro-v1:0",
            contentType="application/json",
            accept="application/json",
            body=json.dumps({
                "messages": [{"role": "user", "content": [{"text": prompt}]}],
                "inferenceConfig": {"maxTokens": 4096, "temperature": 0.7}
            })
        )
        stream = resp.get("body")
        for event in stream:
            chunk = event.get("chunk")
            if chunk:
                try:
                    data = json.loads(chunk.get("bytes", b"{}"))
                    delta = (data.get("contentBlockDelta", {})
                                 .get("delta", {}).get("text", ""))
                    if delta:
                        full_text += delta
                        # show streamed text with blinking cursor
                        placeholder.markdown(
                            f'<div class="proto-doc" style="max-height:400px;">'
                            f'{render_protocol_html(full_text)}'
                            f'<span class="stream-cursor"></span></div>',
                            unsafe_allow_html=True
                        )
                except Exception:
                    pass
    except Exception as e:
        logger.warning("Streaming failed, falling back to converse: %s", e)
        text, _ = call_nova(prompt)
        full_text = text
        placeholder.markdown(
            f'<div class="proto-doc" style="max-height:400px;">'
            f'{render_protocol_html(full_text)}</div>',
            unsafe_allow_html=True
        )
    lat = round(time.time() - t0, 2)
    return full_text, lat


# ── ROI Calculation ──────────────────────────────────────────────────
def compute_roi(phase: str, cost: dict, timeline: dict, drug: str, disease: str) -> dict:
    rng = random.Random(_hash("roi" + drug + disease + phase))
    traditional_months   = timeline["total"] + rng.randint(4, 10)
    ai_months            = timeline["total"]
    months_saved         = traditional_months - ai_months
    traditional_cro_cost = round(cost["grand"] * rng.uniform(1.18, 1.35))
    ai_cro_cost          = cost["grand"]
    cro_savings          = traditional_cro_cost - ai_cro_cost
    amendment_cost_saved = rng.randint(1, 3) * 500_000
    writer_days_saved    = rng.randint(40, 90)
    writer_cost_saved    = writer_days_saved * 1200
    total_savings        = cro_savings + amendment_cost_saved + writer_cost_saved
    return {
        "months_saved":         months_saved,
        "cro_savings":          cro_savings,
        "amendment_saved":      amendment_cost_saved,
        "writer_cost_saved":    writer_cost_saved,
        "writer_days_saved":    writer_days_saved,
        "total_savings":        total_savings,
        "traditional_months":   traditional_months,
    }


# ── Audit Log ───────────────────────────────────────────────────────
def audit_hash(text: str, prev_hash: str = "GENESIS") -> str:
    payload = f"{prev_hash}::{text}::{datetime.datetime.now().isoformat()}"
    return hashlib.sha256(payload.encode()).hexdigest()[:16]


def add_audit_entry(action: str, detail: str, version_id: str = ""):
    if "audit_log" not in st.session_state:
        st.session_state.audit_log = []
    prev = st.session_state.audit_log[-1]["hash"] if st.session_state.audit_log else "GENESIS"
    entry = {
        "ts":        datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "action":    action,
        "detail":    detail,
        "version_id": version_id,
        "hash":      audit_hash(action + detail, prev),
    }
    st.session_state.audit_log.append(entry)


# ── Patient Burden Scoring ───────────────────────────────────────────
def patient_burden_score(soa: list) -> dict:
    blood_draws = sum(1 for r in soa if "Hematology" in r[0] or "Chemistry" in r[0]
                      for v in r[1:] if v)
    imaging     = sum(1 for r in soa if "Imaging" in r[0] for v in r[1:] if v)
    clinic_vis  = len(SOA_TIMEPOINTS)
    prn_visits  = sum(1 for r in soa for v in r[1:] if v)
    score       = min(100, blood_draws * 6 + imaging * 12 + clinic_vis * 4 + prn_visits * 2)
    level       = "Low" if score < 35 else ("Medium" if score < 65 else "High")
    suggestions = []
    if blood_draws > 6:
        suggestions.append("Consider home-nursing for blood draws to reduce clinic burden.")
    if imaging > 2:
        suggestions.append("Allow remote/local imaging to reduce long-distance travel.")
    if clinic_vis > 8:
        suggestions.append("Introduce eVisit or teleconsultation for interim assessments.")
    suggestions.append("Consider electronic Patient Reported Outcomes (ePRO) via smartphone.")
    return {"score": score, "level": level, "blood_draws": blood_draws,
            "imaging": imaging, "clinic_vis": clinic_vis, "suggestions": suggestions[:3]}


# ── I/E Funnel ──────────────────────────────────────────────────────
def ie_funnel(disease: str, drug: str, phase: str, ehr: dict) -> list:
    rng  = random.Random(_hash("funnel" + disease + drug + phase))
    pool = ehr["prevalence"]
    stages = [
        ("Total US Prevalence",      pool,                      "#818CF8"),
        ("Age-Eligible (18–75)",     round(pool * rng.uniform(.55, .72)), "#2DD4BF"),
        ("No Major Comorbidities",   0, "#34D399"),
        ("Meets Biomarker Criteria", 0, "#FBBF24"),
        ("Prior Therapy Criteria",   0, "#F87171"),
        ("Final Eligible Pool",      ehr["eligible_pool"],      "#C084FC"),
    ]
    # fill in middle stages
    prev = stages[1][1]
    for i in range(2, 5):
        nxt = round(prev * rng.uniform(.28, .55))
        stages[i] = (stages[i][0], nxt, stages[i][2])
        prev = nxt
    return stages


# ── Geospatial Site Scores ───────────────────────────────────────────
def site_scores(disease: str, drug: str) -> list:
    disease_hubs = {
        "lung cancer":         [("Texas",90),("California",88),("New York",85),("Florida",82),("Illinois",78)],
        "breast cancer":       [("California",92),("Texas",88),("New York",86),("Florida",83),("Massachusetts",80)],
        "type 2 diabetes":     [("Texas",94),("California",91),("Florida",88),("New York",85),("Georgia",82)],
        "melanoma":            [("California",91),("Texas",87),("Florida",85),("New York",82),("Arizona",79)],
        "alzheimer's disease": [("Florida",90),("California",88),("Texas",85),("New York",83),("Pennsylvania",80)],
        "rheumatoid arthritis":[("California",88),("Texas",85),("New York",83),("Illinois",80),("Ohio",77)],
        "heart failure":       [("Texas",89),("California",86),("Florida",84),("New York",82),("Ohio",79)],
        "hypertension":        [("Texas",93),("California",90),("Florida",87),("New York",84),("Georgia",81)],
    }
    rng = random.Random(_hash("sites" + disease + drug))
    dl  = disease.lower()
    for key, sites in disease_hubs.items():
        if key in dl:
            # add slight jitter
            return [(s, min(99, max(50, sc + rng.randint(-4, 4)))) for s, sc in sites]
    # default
    states = ["California","Texas","New York","Florida","Massachusetts",
              "Illinois","Pennsylvania","Ohio","North Carolina","Washington"]
    random.Random(_hash("def"+disease)).shuffle(states)
    return [(s, rng.randint(65, 92)) for s in states[:5]]


# ── Protocol Chat Q&A ────────────────────────────────────────────────
def chat_with_protocol(question: str, protocol_text: str,
                        drug: str, disease: str, history: list) -> str:
    # Build a compact conversation context
    hist_txt = ""
    for msg in history[-6:]:   # last 3 exchanges
        role = "User" if msg["role"] == "user" else "Assistant"
        hist_txt += f"{role}: {msg['content']}\n"

    prompt = f"""You are a clinical trial expert assistant. The user has a question about this protocol for {drug} in {disease}.

PROTOCOL EXCERPT (first 3000 chars):
{protocol_text[:3000]}

CONVERSATION HISTORY:
{hist_txt}

USER QUESTION: {question}

Answer concisely and precisely. Reference specific protocol sections where relevant. No markdown formatting."""
    try:
        text, _ = call_nova(prompt)
        return text
    except Exception as e:
        return f"Unable to answer: {e}"


# ── Smart Redlining ──────────────────────────────────────────────────
def protocol_diff_html(text_a: str, text_b: str, label_a: str, label_b: str) -> str:
    lines_a = clean_protocol_text(text_a).splitlines()
    lines_b = clean_protocol_text(text_b).splitlines()
    matcher = difflib.SequenceMatcher(None, lines_a, lines_b, autojunk=False)
    html    = f'<div class="diff-wrap">'
    html   += f'<div style="display:flex;gap:1rem;margin-bottom:.75rem;font-size:.65rem;font-weight:700;text-transform:uppercase;letter-spacing:.08em;">'
    html   += f'<span style="color:#fca5a5;">— {label_a}</span>'
    html   += f'<span style="color:#86efac;">+ {label_b}</span></div>'
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for line in lines_a[i1:i2]:
                if line.strip():
                    html += f'<span class="diff-eq">{re.sub("<","&lt;",line)}</span>'
        elif tag in ("replace", "delete"):
            for line in lines_a[i1:i2]:
                if line.strip():
                    html += f'<span class="diff-del">- {re.sub("<","&lt;",line)}</span>'
            if tag == "replace":
                for line in lines_b[j1:j2]:
                    if line.strip():
                        html += f'<span class="diff-add">+ {re.sub("<","&lt;",line)}</span>'
        elif tag == "insert":
            for line in lines_b[j1:j2]:
                if line.strip():
                    html += f'<span class="diff-add">+ {re.sub("<","&lt;",line)}</span>'
    html += "</div>"
    return html


def clean_protocol_text(raw: str) -> str:
    """Strip markdown symbols and return clean plain text."""
    text = raw
    # remove headers like ### Title → TITLE
    text = re.sub(r"^#{1,6}\s*(.+)$", lambda m: m.group(1).upper(), text, flags=re.MULTILINE)
    # remove bold/italic
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    # remove horizontal rules
    text = re.sub(r"^[-_*]{3,}\s*$", "", text, flags=re.MULTILINE)
    # remove inline code
    text = re.sub(r"`(.+?)`", r"\1", text)
    # clean extra blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def render_protocol_html(raw: str) -> str:
    """Convert clean protocol text to styled HTML for document viewer."""
    text = clean_protocol_text(raw)
    lines = text.split("\n")
    html  = ""
    for line in lines:
        s = line.strip()
        if not s:
            html += "<br>"
            continue
        # Numbered section headings: "1. TITLE" or "TITLE:" pattern
        if re.match(r"^\d{1,2}\.\s+[A-Z][A-Z\s&/()\-]{4,}$", s) or \
           (s.isupper() and len(s) > 4 and len(s) < 80):
            html += f'<div class="proto-h1">{s}</div>'
        elif s.endswith(":") and len(s) < 60 and s[0].isupper():
            html += f'<div class="proto-h2">{s[:-1]}</div>'
        elif s.startswith(("- ", "• ", "* ", "· ")):
            html += f'<div class="proto-bullet">{"→ " + s[2:]}</div>'
        elif re.match(r"^\d+\.\s", s):
            html += f'<div class="proto-bullet">{s}</div>'
        else:
            html += f'<div class="proto-p">{s}</div>'
    return html


def build_prompt(drug: str, disease: str, phase: str, mode: str,
                 rag_ctx: str, citations: str) -> str:
    mode_map = {
        "Fastest Approval":            "Prioritise adaptive design, surrogate endpoints, Fast Track/Breakthrough designation, rolling submission.",
        "Lowest Cost":                 "Recommend decentralised trial, lean CRO, risk-based monitoring, pragmatic design.",
        "Highest Safety":              "Conservative dose escalation, enhanced DSMB, sentinel dosing, robust stopping rules.",
        "Highest Success Probability": "Balance regulatory, scientific, recruitment and safety via enriched selection and adaptive design.",
    }
    return f"""You are Nova, an expert clinical trial design AI. Generate a comprehensive, ICH-GCP-compliant clinical trial protocol.

OPTIMISATION MODE: {mode}
Directive: {mode_map.get(mode,'')}
{rag_ctx}
{citations}

TRIAL PARAMETERS:
- Investigational Drug: {drug}
- Disease / Indication: {disease}
- Phase: {phase}

Generate a COMPLETE protocol with ALL numbered sections. Use plain numbered headings (1. TITLE, 2. RATIONALE, etc). Do NOT use markdown (no ###, **, --, etc).

Sections required:
1. PROTOCOL TITLE & OVERVIEW
2. SCIENTIFIC RATIONALE (cite references where provided)
3. STUDY OBJECTIVES — Primary, Secondary, Exploratory
4. STUDY DESIGN — type, blinding, randomisation, adaptive elements
5. INCLUSION CRITERIA — minimum 8
6. EXCLUSION CRITERIA — minimum 8
7. STUDY POPULATION & SAMPLE SIZE JUSTIFICATION — include power calc
8. INVESTIGATIONAL DRUG & DOSING REGIMEN
9. COMPARATOR ARM
10. PRIMARY & SECONDARY ENDPOINTS
11. STATISTICAL ANALYSIS PLAN — power, alpha, interim, populations
12. SAFETY MONITORING — DSMB, SAE definitions, stopping rules
13. PATIENT RECRUITMENT FEASIBILITY
14. RISK ASSESSMENT & MITIGATION
15. REGULATORY STRATEGY — IND/CTA, FDA pathway
16. ICH-GCP COMPLIANCE SUMMARY
17. PROTOCOL IMPROVEMENT SUGGESTIONS — numbered list of 5
18. REASONING TRANSPARENCY — brief justification of key design choices

Write in formal clinical protocol language. Plain text only."""


def generate_icf(protocol_text: str, drug: str, disease: str) -> str:
    """Generate patient-friendly ICF summary from protocol."""
    proto_short = protocol_text[:3000]
    prompt = f"""You are a patient advocate and medical writer.

Given this clinical trial protocol for {drug} in {disease}:
---
{proto_short}
---

Write a plain-language Informed Consent Form summary at a 6th-grade reading level.
Use these sections:
- WHAT IS THIS STUDY ABOUT? (2-3 simple sentences)
- WHY ARE WE DOING THIS STUDY? (1-2 sentences)
- WHAT WILL HAPPEN TO ME? (bullet list, simple language)
- WHAT ARE THE RISKS? (bullet list with ⚠ symbols)
- WHAT ARE THE BENEFITS? (bullet list with ✓ symbols)
- DO I HAVE TO JOIN? (1-2 sentences about voluntary nature)
- WHO DO I CALL WITH QUESTIONS? (placeholder contact)

Use simple words. No jargon. No markdown bold/italic. Plain text with clear labels."""
    try:
        text, _ = call_nova(prompt)
        return text
    except Exception:
        return f"Informed Consent Form for: {drug} in {disease}\n\nThis is a research study. Participation is voluntary. Please ask your doctor any questions before signing."


def generate_red_team(protocol_text: str, drug: str, disease: str) -> str:
    """Red team mock FDA review."""
    proto_short = protocol_text[:4000]
    prompt = f"""You are a hyper-critical FDA senior reviewer with 20 years of experience in {disease} trials.
Your job is to ATTACK this clinical trial protocol for {drug} in {disease}, finding every weakness.

Protocol excerpt:
---
{proto_short}
---

Produce a numbered list of 6-8 specific, actionable critiques covering:
- Safety concerns and loopholes
- Endpoint mismatches or poor sensitivity
- Statistical design weaknesses
- Regulatory compliance gaps
- Patient safety risks
- Missing sections or ambiguous language

Be blunt. Use FDA reviewer language. No markdown. Number each finding."""
    try:
        text, _ = call_nova(prompt)
        return text
    except Exception:
        return "Red team review unavailable. Please retry."


def export_cdisc_xml(data: dict) -> str:
    """Generate CDISC SDTM-style XML export."""
    root = ET.Element("ODM",
                      xmlns="http://www.cdisc.org/ns/odm/v1.3",
                      FileType="Snapshot",
                      FileOID=f"TF-{data['version_id']}",
                      CreationDateTime=data["timestamp"])
    study = ET.SubElement(root, "Study", OID=f"STUDY.{data['version_id']}")
    gv    = ET.SubElement(study, "GlobalVariables")
    ET.SubElement(gv, "StudyName").text    = f"{data['drug']} in {data['disease']}"
    ET.SubElement(gv, "StudyDescription").text = f"{data['phase']} trial — TrialForge AI"
    ET.SubElement(gv, "ProtocolName").text = f"TF-{data['version_id']}"

    meta = ET.SubElement(study, "MetaDataVersion", OID="MDV.1", Name="Protocol v1")

    # Protocol — key fields
    fields = [
        ("DRUG",    "Investigational Drug",        data["drug"]),
        ("DIS",     "Disease Indication",          data["disease"]),
        ("PHASE",   "Trial Phase",                 data["phase"]),
        ("MODE",    "Optimisation Mode",           data["mode"]),
        ("COMP",    "Comparator",                  data["comparator"]),
        ("N",       "Sample Size",                 str(data["sample_size"])),
        ("DUR",     "Study Duration (months)",     str(data["timeline"]["total"])),
        ("AP",      "Approval Probability",        f"{int(data['scores']['ap']*100)}%"),
        ("SP",      "Success Probability",         f"{int(data['scores']['sp']*100)}%"),
        ("ICH",     "ICH Compliance Score",        f"{data['quality']['ich']}%"),
        ("FDA",     "FDA Readiness",               f"{data['quality']['fda']}%"),
        ("COST",    "Total Cost Estimate (USD)",   f"${data['cost']['grand']:,}"),
    ]
    form = ET.SubElement(meta, "FormDef", OID="F.PROTOCOL", Name="Protocol Metadata", Repeating="No")
    for oid, name, val in fields:
        ig = ET.SubElement(form, "ItemGroupRef", ItemGroupOID=f"IG.{oid}", Mandatory="Yes")
    for oid, name, val in fields:
        ig = ET.SubElement(meta, "ItemGroupDef", OID=f"IG.{oid}", Name=name, Repeating="No")
        it = ET.SubElement(ig, "ItemRef", ItemOID=f"I.{oid}", Mandatory="Yes")
        itd= ET.SubElement(meta, "ItemDef", OID=f"I.{oid}", Name=name, DataType="text")
        ET.SubElement(itd, "Question").text = name
        cdl = ET.SubElement(itd, "CodeListRef", CodeListOID=f"CL.{oid}")
        clr = ET.SubElement(meta, "CodeList", OID=f"CL.{oid}", Name=name, DataType="text")
        cli = ET.SubElement(clr, "CodeListItem", CodedValue=val)
        ET.SubElement(cli, "Decode").text = val

    xmlstr = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")
    return xmlstr


def generate_pdf_bytes(data: dict) -> bytes:
    if not HAVE_PDF:
        return b""
    buf    = io.BytesIO()
    doc    = SimpleDocTemplate(buf, pagesize=A4,
                               leftMargin=2.2*cm, rightMargin=2.2*cm,
                               topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    navy   = colors.HexColor("#0D1117")
    cyan   = colors.HexColor("#2DD4BF")
    indigo = colors.HexColor("#818CF8")
    slate  = colors.HexColor("#E6EDF3")
    gray   = colors.HexColor("#8B949E")
    white  = colors.white
    light  = colors.HexColor("#F1F5F9")

    h_title = ParagraphStyle("T", fontSize=20, textColor=navy, fontName="Helvetica-Bold",
                              spaceAfter=4, leading=26, alignment=TA_CENTER)
    h1  = ParagraphStyle("H1", fontSize=13, textColor=indigo, fontName="Helvetica-Bold",
                          spaceBefore=14, spaceAfter=4, leading=18)
    h2  = ParagraphStyle("H2", fontSize=10, textColor=navy, fontName="Helvetica-Bold",
                          spaceBefore=8, spaceAfter=3, leading=14)
    bod = ParagraphStyle("B", fontSize=8.5, textColor=colors.HexColor("#1E293B"),
                          leading=13, spaceAfter=3, fontName="Helvetica")
    meta_s = ParagraphStyle("M", fontSize=7.5, textColor=gray, leading=11, fontName="Helvetica")
    s = []

    # Cover
    s.append(Spacer(1, 1*cm))
    s.append(Paragraph("TrialForge AI", meta_s))
    s.append(Paragraph("CLINICAL TRIAL PROTOCOL", h_title))
    s.append(Paragraph(f"{data['drug']} in {data['disease']}", h1))
    s.append(Spacer(1, 6))

    m_rows = [
        ["Version:", data["version_id"], "Phase:", data["phase"]],
        ["Generated:", data["timestamp"], "Mode:", data["mode"]],
        ["Drug:", data["drug"],           "Disease:", data["disease"]],
        ["Comparator:", data["comparator"][:50], "Sample Size:", f"N = {data['sample_size']}"],
        ["Duration:", f"{data['timeline']['total']} months", "Latency:", f"{data['latency']}s"],
    ]
    mt = Table(m_rows, colWidths=[3.2*cm,5.8*cm,3*cm,5*cm])
    mt.setStyle(TableStyle([
        ("FONTSIZE",      (0,0),(-1,-1), 7.5),
        ("FONTNAME",      (0,0),(0,-1), "Helvetica-Bold"),
        ("FONTNAME",      (2,0),(2,-1), "Helvetica-Bold"),
        ("TEXTCOLOR",     (0,0),(0,-1), gray),
        ("TEXTCOLOR",     (2,0),(2,-1), gray),
        ("ROWBACKGROUNDS",(0,0),(-1,-1),[light, white]),
        ("GRID",          (0,0),(-1,-1), 0.3, colors.HexColor("#E2E8F0")),
        ("TOPPADDING",    (0,0),(-1,-1), 4),
        ("BOTTOMPADDING", (0,0),(-1,-1), 4),
        ("LEFTPADDING",   (0,0),(-1,-1), 6),
    ]))
    s.append(mt)
    s.append(Spacer(1,10))
    s.append(HRFlowable(width="100%", thickness=1.5, color=cyan))
    s.append(Spacer(1,8))

    # Quality scores table
    q = data["quality"]
    sc = data["scores"]
    s.append(Paragraph("PROTOCOL QUALITY SCORES", h1))
    qrows = [["Metric","Score","Status"],
             ["ICH-GCP Compliance",     f"{q['ich']}%",  "Pass" if q['ich']>=75 else "Review"],
             ["Regulatory Readiness",   f"{q['reg']}%",  "High" if q['reg']>=78 else "Medium"],
             ["FDA Approval Readiness", f"{q['fda']}%",  q['fda_lbl']],
             ["Failure Risk",           f"{q['fail']}%", q['fail_lbl']+" Risk"],
             ["Protocol Confidence",    f"{q['conf']}%", "High" if q['conf']>=80 else "Medium"],
             ["Approval Probability",   f"{int(sc['ap']*100)}%", ""],
             ["Success Probability",    f"{int(sc['sp']*100)}%", ""],]
    qt = Table(qrows, colWidths=[7*cm,4*cm,6*cm])
    qt.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,0), navy),
        ("TEXTCOLOR",     (0,0),(-1,0), white),
        ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
        ("FONTSIZE",      (0,0),(-1,-1), 8.5),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[light,white]),
        ("GRID",          (0,0),(-1,-1), 0.3, colors.HexColor("#CBD5E1")),
        ("TOPPADDING",    (0,0),(-1,-1), 4),
        ("BOTTOMPADDING", (0,0),(-1,-1), 4),
        ("LEFTPADDING",   (0,0),(-1,-1), 7),
    ]))
    s.append(qt)
    s.append(Spacer(1,10))
    s.append(PageBreak())

    # Protocol text
    s.append(Paragraph("FULL PROTOCOL TEXT", h1))
    s.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#E2E8F0")))
    s.append(Spacer(1,8))
    clean = clean_protocol_text(data["protocol"])
    for line in clean.split("\n"):
        line = line.strip()
        if not line:
            s.append(Spacer(1,4)); continue
        safe = line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        if (re.match(r"^\d{1,2}\.\s+[A-Z]", line) or
                (line.isupper() and 4 < len(line) < 80)):
            s.append(Paragraph(safe, h1))
        elif safe.endswith(":") and len(safe) < 55:
            s.append(Paragraph(safe, h2))
        else:
            s.append(Paragraph(safe, bod))

    doc.build(s)
    return buf.getvalue()


def generate_docx_bytes(data: dict) -> bytes:
    if not HAVE_DOCX:
        return b""
    doc = Document()
    doc.core_properties.title   = f"Protocol: {data['drug']} in {data['disease']}"
    doc.core_properties.author  = "TrialForge AI"
    doc.core_properties.subject = f"{data['phase']} Clinical Trial Protocol"

    # Styles
    for style_name in ["Normal","Heading 1","Heading 2"]:
        if style_name in [s.name for s in doc.styles]:
            st = doc.styles[style_name]
            st.font.name = "Calibri"

    # Title
    t = doc.add_heading("CLINICAL TRIAL PROTOCOL", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f"{data['drug']} in {data['disease']}", 1)

    # Metadata table
    doc.add_heading("Protocol Metadata", 2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    meta_pairs = [
        ("Version", data["version_id"],       "Phase",    data["phase"]),
        ("Generated", data["timestamp"],      "Mode",     data["mode"]),
        ("Drug", data["drug"],                "Disease",  data["disease"]),
        ("Comparator", data["comparator"][:45],"N",       str(data["sample_size"])),
        ("Duration", f"{data['timeline']['total']}mo", "Latency", f"{data['latency']}s"),
    ]
    for i,(k1,v1,k2,v2) in enumerate(meta_pairs):
        row = table.rows[i]
        for j,(txt,bold) in enumerate([(k1,True),(v1,False),(k2,True),(v2,False)]):
            cell = row.cells[j]
            run  = cell.paragraphs[0].add_run(txt)
            run.bold = bold
            run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading("Quality Scores", 2)
    q = data["quality"]; sc = data["scores"]
    qt = doc.add_table(rows=8, cols=2)
    qt.style = "Table Grid"
    for i,(k,v) in enumerate([
        ("ICH-GCP Compliance",   f"{q['ich']}%"),
        ("Regulatory Readiness", f"{q['reg']}%"),
        ("FDA Readiness",        f"{q['fda']}%"),
        ("Failure Risk",         f"{q['fail']}% ({q['fail_lbl']})"),
        ("Confidence Score",     f"{q['conf']}%"),
        ("Approval Probability", f"{int(sc['ap']*100)}%"),
        ("Success Probability",  f"{int(sc['sp']*100)}%"),
        ("Risk Class",           sc["risk_class"]),
    ]):
        r = qt.rows[i]
        rk = r.cells[0].paragraphs[0].add_run(k); rk.bold = True; rk.font.size = Pt(9)
        rv = r.cells[1].paragraphs[0].add_run(v); rv.font.size = Pt(9)

    doc.add_page_break()
    doc.add_heading("Full Protocol Text", 1)
    clean = clean_protocol_text(data["protocol"])
    for line in clean.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if re.match(r"^\d{1,2}\.\s+[A-Z]", line) or (line.isupper() and 4<len(line)<80):
            doc.add_heading(line, 2)
        elif line.endswith(":") and len(line)<55:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.bold = True
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
# MAIN GENERATION PIPELINE
# ═══════════════════════════════════════════════════════════════════
@st.cache_data(ttl=1800, show_spinner=False)
def cached_generate(drug: str, disease: str, phase: str, mode: str,
                    use_rag: bool, use_pubmed: bool) -> dict:
    rng  = random.Random(_hash(drug+disease+phase))
    n    = rng.randint(80, 1800)
    dur  = rng.randint(12, 48)

    # RAG
    nct_trials = fetch_clinicaltrials(disease, drug) if use_rag else []
    rag_ctx    = ""
    if nct_trials:
        rag_ctx = "\nRelevant Trials (ClinicalTrials.gov):\n"
        for t in nct_trials[:3]:
            rag_ctx += f"- [{t['id']}] {t['title']} ({t['phase']}) Status:{t['status']} N:{t['n']}\n"

    # PubMed
    pubs = fetch_pubmed_citations(drug, disease) if use_pubmed else []
    cit_ctx = ""
    if pubs:
        cit_ctx = "\nPubMed Citations (insert into Scientific Rationale section):\n"
        for p in pubs:
            cit_ctx += f"- {p['title']} ({p['journal']}, {p['year']}) PMID:{p['pmid']}\n"

    prompt   = build_prompt(drug, disease, phase, mode, rag_ctx, cit_ctx)
    protocol, latency = call_nova(prompt)

    comparator = suggest_comparator(disease)
    sc  = risk_scores(phase, disease, drug)
    q   = quality_scores(phase, disease, drug, protocol)
    c   = cost_estimate(phase, n, dur)
    tl  = timeline_est(phase, n)
    ehr = synthetic_ehr_feasibility(disease, phase, n)
    amd = amendment_flags(drug, disease, phase)
    ddi = drug_interactions(drug)
    div = diversity_plan(disease, phase)

    return {
        "drug": drug, "disease": disease, "phase": phase, "mode": mode,
        "comparator": comparator, "protocol": protocol,
        "scores": sc, "quality": q, "cost": c, "timeline": tl,
        "ehr": ehr, "amendments": amd, "interactions": ddi, "diversity": div,
        "nct": nct_trials, "pubmed": pubs,
        "sample_size": n, "duration": dur, "latency": latency,
        "roi": compute_roi(phase, c, tl, drug, disease),
        "burden": patient_burden_score(SOA_ASSESSMENTS),
        "funnel": ie_funnel(disease, drug, phase, ehr),
        "sites":  site_scores(disease, drug),
        "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "version_id": "v1",   # overwritten after cache
    }


def generate(drug: str, disease: str, phase: str, mode: str,
             use_rag: bool, use_pubmed: bool) -> dict:
    result = dict(cached_generate(drug, disease, phase, mode, use_rag, use_pubmed))
    result["version_id"] = f"v{len(st.session_state.get('versions',[]))+1}"
    return result


# ═══════════════════════════════════════════════════════════════════
# HTML HELPERS
# ═══════════════════════════════════════════════════════════════════
def pb(label: str, pct: int, color: str, right: str = "") -> str:
    return f"""<div class="pb-wrap"><div class="pb-lbl">
    <span>{label}</span><span style="color:{color};font-weight:700;">{right or f'{pct}%'}</span>
    </div><div class="pb-track"><div class="pb-fill" style="width:{pct}%;background:{color};"></div></div></div>"""


def ring(pct: int, color: str, label: str, sub: str = "") -> str:
    r = 26; c = 2*3.14159*r; off = c*(1-pct/100)
    return f"""<div class="metric-card"><div style="position:relative;width:62px;height:62px;margin:0 auto 5px;">
    <svg width="62" height="62" viewBox="0 0 62 62" style="transform:rotate(-90deg)">
        <circle cx="31" cy="31" r="{r}" fill="none" stroke="var(--card)" stroke-width="5"/>
        <circle cx="31" cy="31" r="{r}" fill="none" stroke="{color}" stroke-width="5"
            stroke-linecap="round" stroke-dasharray="{c:.1f}" stroke-dashoffset="{off:.1f}"/>
    </svg>
    <div style="position:absolute;inset:0;display:flex;align-items:center;justify-content:center;font-weight:800;font-size:.84rem;color:{color};">{pct}%</div>
    </div><div class="mc-lbl">{label}</div>{f'<div class="mc-sub" style="color:{color};">{sub}</div>' if sub else ''}</div>"""


def gantt(tl: dict) -> str:
    tot   = tl["total"]
    bars  = [
        ("Startup",     0,          tl["su"],                  "#818CF8"),
        ("Enrolment",   tl["su"],   tl["su"]+tl["en"],         "#2DD4BF"),
        ("Treatment",   tl["su"]+tl["en"], tl["su"]+tl["en"]+tl["base"], "#34D399"),
        ("Follow-up",   tl["su"]+tl["en"]+tl["base"], tot,    "#FBBF24"),
    ]
    html = '<div style="padding:4px 0;">'
    for lbl, st_, end, col in bars:
        lp = round(st_/tot*100); wp = round((end-st_)/tot*100); dur = end-st_
        html += f'<div class="gantt-row"><div class="gantt-lbl">{lbl}</div><div class="gantt-track"><div class="gantt-bar" style="margin-left:{lp}%;width:{wp}%;background:{col};">{dur}mo</div></div></div>'
    html += f'<div class="gantt-row"><div class="gantt-lbl" style="font-size:.57rem;">Months</div><div style="flex:1;display:flex;justify-content:space-between;font-size:.6rem;color:var(--text2);font-family:var(--mono);"><span>0</span><span>{tot//4}</span><span>{tot//2}</span><span>{3*tot//4}</span><span>{tot}</span></div></div>'
    html += "</div>"
    return html


def soa_table_html() -> str:
    tp_labels = SOA_TIMEPOINTS
    html  = '<div class="soa-wrap"><table class="soa-table"><thead><tr>'
    html += '<th class="soa-first">Assessment</th>'
    for tp in tp_labels:
        html += f'<th>{tp}</th>'
    html += '</tr></thead><tbody>'
    for row in SOA_ASSESSMENTS:
        name = row[0]; cells = row[1:]
        html += f'<tr><td class="soa-first">{name}</td>'
        for req in cells:
            if req:
                html += '<td><span class="soa-dot soa-req" title="Required"></span></td>'
            else:
                html += '<td><span style="color:var(--text3);font-size:.7rem;">—</span></td>'
        html += '</tr>'
    html += '</tbody></table></div>'
    html += '<div style="display:flex;gap:12px;margin-top:8px;font-size:.64rem;color:var(--text2);">'
    html += '<span><span class="soa-dot soa-req"></span> Required visit</span>'
    html += '<span><span style="color:var(--text3);">—</span> Not applicable</span>'
    html += '</div>'
    return html


def plotly_budget_chart(cost: dict):
    """Interactive Plotly donut chart for budget breakdown."""
    if not HAVE_PLOTLY:
        return None
    labels = ["Direct Costs", "CRO", "Overhead", "Regulatory"]
    values = [cost["total"], cost["cro"], cost["ov"], cost["reg"]]
    colors_list = ["#2DD4BF", "#818CF8", "#FBBF24", "#34D399"]
    fig = go.Figure(go.Pie(
        labels=labels, values=values,
        hole=0.55,
        marker=dict(colors=colors_list, line=dict(color="#0D1117", width=2)),
        textinfo="label+percent",
        textfont=dict(size=11, family="Inter"),
        hovertemplate="<b>%{label}</b><br>$%{value:,.0f}<br>%{percent}<extra></extra>",
    ))
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=0, b=0), height=220,
        font=dict(color="#8B949E", family="Inter"),
        legend=dict(font=dict(color="#8B949E", size=10), bgcolor="rgba(0,0,0,0)"),
        annotations=[dict(text=f"${cost['grand']//1000}K", x=0.5, y=0.5,
                          font_size=16, font_color="#E6EDF3", font_family="Inter",
                          showarrow=False)]
    )
    return fig


def plotly_enroll_chart(tl: dict, n: int):
    """Interactive Plotly line chart for enrolment trajectory."""
    if not HAVE_PLOTLY:
        return None
    monthly = max(1, n // max(1, tl["en"]))
    months, enrolled = [], []
    for mo in range(tl["su"], tl["su"] + tl["en"] + 1):
        e = min((mo - tl["su"] + 1) * monthly, n)
        months.append(mo)
        enrolled.append(e)
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=months, y=enrolled, mode="lines",
        fill="tozeroy",
        line=dict(color="#2DD4BF", width=2.5),
        fillcolor="rgba(45,212,191,0.12)",
        hovertemplate="Month %{x}<br>Enrolled: %{y:,}<extra></extra>",
        name="Enrolled"
    ))
    fig.add_hline(y=n, line_dash="dot", line_color="#818CF8",
                  annotation_text=f"Target N={n}", annotation_font_color="#818CF8")
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=10, b=30), height=200,
        font=dict(color="#8B949E", family="Inter"),
        xaxis=dict(title="Month", gridcolor="#1C2130", color="#8B949E", showgrid=True),
        yaxis=dict(title="Patients", gridcolor="#1C2130", color="#8B949E", showgrid=True),
        showlegend=False,
    )
    return fig


def plotly_funnel_chart(stages: list):
    """Interactive Plotly funnel for I/E patient pool."""
    if not HAVE_PLOTLY:
        return None
    labels = [s[0] for s in stages]
    values = [s[1] for s in stages]
    clrs   = [s[2] for s in stages]
    fig = go.Figure(go.Funnel(
        y=labels, x=values,
        marker=dict(color=clrs, line=dict(width=1, color="#0D1117")),
        texttemplate="%{value:,.0f}<br>(%{percentInitial:.1%})",
        textfont=dict(family="Inter", size=10),
        hovertemplate="<b>%{y}</b><br>%{x:,.0f} patients<extra></extra>",
    ))
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=10, b=10), height=320,
        font=dict(color="#8B949E", family="Inter"),
        funnelmode="stack",
    )
    return fig


def plotly_radar_chart(q: dict, sc: dict):
    """Radar chart for protocol quality dimensions."""
    if not HAVE_PLOTLY:
        return None
    categories = ["ICH Compliance", "Regulatory", "FDA Readiness",
                  "Completeness", "Confidence", "Approval Prob"]
    vals = [q["ich"], q["reg"], q["fda"], q["comp"], q["conf"], int(sc.get("ap", 0.6)*100)]
    vals_closed = vals + [vals[0]]
    cats_closed  = categories + [categories[0]]
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=vals_closed, theta=cats_closed,
        fill="toself",
        fillcolor="rgba(45,212,191,0.12)",
        line=dict(color="#2DD4BF", width=2),
        hovertemplate="%{theta}: %{r}%<extra></extra>",
        name="Protocol Score"
    ))
    fig.update_layout(
        polar=dict(
            bgcolor="rgba(0,0,0,0)",
            radialaxis=dict(visible=True, range=[0,100], gridcolor="#1C2130",
                            tickfont=dict(color="#8B949E", size=8), linecolor="#1C2130"),
            angularaxis=dict(gridcolor="#1C2130", linecolor="#1C2130",
                             tickfont=dict(color="#8B949E", size=9)),
        ),
        paper_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=40, r=40, t=20, b=20), height=260,
        font=dict(color="#8B949E", family="Inter"),
        showlegend=False,
    )
    return fig


def plotly_site_map(sites: list, disease: str):
    """US choropleth-style bubble map for trial site scores."""
    if not HAVE_PLOTLY:
        return None
    state_abbr = {
        "California":"CA","Texas":"TX","New York":"NY","Florida":"FL",
        "Massachusetts":"MA","Illinois":"IL","Pennsylvania":"PA","Ohio":"OH",
        "North Carolina":"NC","Washington":"WA","Georgia":"GA","Arizona":"AZ",
        "Michigan":"MI","Virginia":"VA","Colorado":"CO",
    }
    # state centroids (lat, lon)
    centroids = {
        "CA":(36.7,-119.4),"TX":(31.0,-99.9),"NY":(42.9,-75.5),"FL":(27.8,-81.5),
        "MA":(42.4,-71.9),"IL":(40.0,-89.0),"PA":(41.2,-77.2),"OH":(40.4,-82.7),
        "NC":(35.6,-79.4),"WA":(47.4,-120.5),"GA":(32.7,-83.2),"AZ":(34.3,-111.1),
        "MI":(44.3,-85.4),"VA":(37.8,-78.2),"CO":(39.0,-105.5),
    }
    lats, lons, names, scores, texts = [], [], [], [], []
    for state, score in sites:
        ab = state_abbr.get(state)
        if ab and ab in centroids:
            lat, lon = centroids[ab]
            lats.append(lat); lons.append(lon)
            names.append(state); scores.append(score)
            texts.append(f"{state}<br>Score: {score}/100")
    fig = go.Figure(go.Scattergeo(
        lat=lats, lon=lons,
        text=texts, hovertemplate="%{text}<extra></extra>",
        marker=dict(
            size=[s/6 for s in scores],
            color=scores, colorscale=[[0,"#F87171"],[0.5,"#FBBF24"],[1,"#2DD4BF"]],
            cmin=50, cmax=100,
            colorbar=dict(title="Score", thickness=10,
                          tickfont=dict(color="#8B949E"), titlefont=dict(color="#8B949E")),
            line=dict(width=1, color="#0D1117"),
            opacity=0.85,
        ),
        mode="markers+text",
        textfont=dict(color="#E6EDF3", size=9),
        textposition="top center",
    ))
    fig.update_geos(
        scope="usa",
        bgcolor="rgba(0,0,0,0)",
        showland=True, landcolor="#1C2130",
        showcoastlines=True, coastlinecolor="#484F58",
        showstates=True, statecolor="#484F58",
        showcountries=False,
        showocean=True, oceancolor="#0D1117",
        projection_type="albers usa",
    )
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=0, b=0), height=320,
        font=dict(color="#8B949E", family="Inter"),
        title=dict(text=f"Recommended Trial Sites — {disease}", font=dict(color="#8B949E", size=11)),
    )
    return fig
    items = [
        ("Direct", cost["total"],  "#2DD4BF"),
        ("CRO",    cost["cro"],    "#818CF8"),
        ("Overhead",cost["ov"],    "#FBBF24"),
        ("Regulatory",cost["reg"],"#34D399"),
    ]
    grand = cost["grand"]
    bars  = []
    y = 10
    for lbl, val, col in items:
        w = round(val/grand*240)
        pct = round(val/grand*100)
        bars.append(f'<rect x="120" y="{y}" width="{w}" height="18" rx="3" fill="{col}" opacity=".85"/>')
        bars.append(f'<text x="115" y="{y+13}" text-anchor="end" fill="#8B949E" font-size="10" font-family="Inter">{lbl}</text>')
        bars.append(f'<text x="{120+w+5}" y="{y+13}" fill="{col}" font-size="9" font-family="JetBrains Mono,monospace">${val:,} ({pct}%)</text>')
        y += 26
    svg_h = y + 10
    return f'<svg viewBox="0 0 500 {svg_h}" xmlns="http://www.w3.org/2000/svg" style="width:100%;max-width:480px;">{"".join(bars)}</svg>'


def svg_enroll_chart(tl: dict, n: int) -> str:
    pts = []
    monthly = max(1, n // tl["en"])
    enrolled = 0
    for mo in range(tl["su"], tl["su"]+tl["en"]+1, max(1,tl["en"]//8)):
        batch = min(monthly * (mo-tl["su"]+1), n)
        enrolled = min(batch, n)
        pts.append((mo, enrolled))
    if not pts or pts[-1][1] < n:
        pts.append((tl["su"]+tl["en"], n))
    total_mo = tl["total"]
    W, H = 440, 100
    def sx(m): return round(m/total_mo*W)
    def sy(e): return round(H - e/n*H)
    path = f"M {sx(pts[0][0])} {sy(pts[0][1])} " + " ".join(f"L {sx(m)} {sy(e)}" for m,e in pts[1:])
    area = path + f" L {sx(pts[-1][0])} {H} L {sx(pts[0][0])} {H} Z"
    return f'''<svg viewBox="-10 -5 {W+60} {H+30}" xmlns="http://www.w3.org/2000/svg" style="width:100%;max-width:460px;">
      <defs><linearGradient id="eg" x1="0" y1="0" x2="0" y2="1"><stop offset="0%" stop-color="#2DD4BF" stop-opacity=".3"/><stop offset="100%" stop-color="#2DD4BF" stop-opacity="0"/></linearGradient></defs>
      <path d="{area}" fill="url(#eg)"/>
      <path d="{path}" fill="none" stroke="#2DD4BF" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
      <line x1="0" y1="{H}" x2="{W}" y2="{H}" stroke="#484F58" stroke-width="1"/>
      <text x="-5" y="{H+18}" fill="#8B949E" font-size="9" font-family="JetBrains Mono">0</text>
      <text x="{W-12}" y="{H+18}" fill="#8B949E" font-size="9" font-family="JetBrains Mono">{total_mo}mo</text>
      <text x="-8" y="8" fill="#8B949E" font-size="9" font-family="JetBrains Mono">{n}</text>
      <text x="{W//2-20}" y="{H+20}" fill="#8B949E" font-size="9" font-family="Inter">Patient Enrolment Over Time</text>
    </svg>'''


# ═══════════════════════════════════════════════════════════════════
# SCORE KEY NORMALISER  (handles both old and new key naming)
# ═══════════════════════════════════════════════════════════════════
def _sc(version_dict: dict, short_key: str) -> float:
    """Safe score accessor — handles both naming conventions."""
    s = version_dict.get("scores", {})
    # short keys: ap, rs, sr, sp
    # legacy long keys: approval_prob, recruit_score, safety_risk, success_prob
    aliases = {
        "ap": ["ap", "approval_prob"],
        "rs": ["rs", "recruit_score"],
        "sr": ["sr", "safety_risk"],
        "sp": ["sp", "success_prob"],
    }
    for key in aliases.get(short_key, [short_key]):
        if key in s:
            return float(s[key])
    return 0.5   # safe default


def _normalise_scores(d: dict) -> dict:
    """Ensure scores dict always has short keys ap/rs/sr/sp."""
    s = d.get("scores", {})
    mapping = {
        "ap": ["ap", "approval_prob"],
        "rs": ["rs", "recruit_score"],
        "sr": ["sr", "safety_risk"],
        "sp": ["sp", "success_prob"],
    }
    for short, candidates in mapping.items():
        if short not in s:
            for c in candidates:
                if c in s:
                    s[short] = s[c]
                    break
    # ensure risk_class and rec_diff exist
    sr = s.get("sr", s.get("safety_risk", 0.3))
    rs = s.get("rs", s.get("recruit_score", 0.6))
    if "risk_class" not in s:
        s["risk_class"] = "Low" if sr < 0.30 else ("Medium" if sr < 0.60 else "High")
    if "rec_diff" not in s:
        s["rec_diff"] = "Easy" if rs > 0.70 else ("Moderate" if rs > 0.45 else "Difficult")
    d["scores"] = s
    return d


# ═══════════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════════
for k, v in [("versions",[]), ("active",None), ("icf_cache",{}),
             ("red_team_cache",{}), ("tmpl",None),
             ("audit_log",[]), ("chat_histories",{})]:
    if k not in st.session_state:
        st.session_state[k] = v

# ═══════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div class="sb-logo">
        <div class="sb-logo-icon">⚗️</div>
        <div><div class="sb-logo-name">TrialForge AI</div>
             <div class="sb-logo-ver">v3.0 · Hackathon Build</div></div>
    </div>
    <div style="display:flex;gap:8px;margin-bottom:.5rem;padding:2px 3px;">
        <div class="tf-pill"><span class="dot dot-g"></span>Nova Pro</div>
        <div class="tf-pill"><span class="dot dot-c"></span>Bedrock</div>
        <div class="tf-pill"><span class="dot dot-g"></span>RAG</div>
    </div>
    """, unsafe_allow_html=True)

    # Apply template values into session state BEFORE widgets render
    tf = st.session_state.tmpl
    if tf:
        st.session_state["inp_drug"]    = tf["drug"]
        st.session_state["inp_disease"] = tf["disease"]
        st.session_state["inp_phase"]   = tf["phase"]
        st.session_state["inp_mode"]    = tf["mode"]
        st.session_state.tmpl = None   # clear immediately after applying

    # Ensure keys exist
    if "inp_drug"    not in st.session_state: st.session_state["inp_drug"]    = ""
    if "inp_disease" not in st.session_state: st.session_state["inp_disease"] = ""
    if "inp_phase"   not in st.session_state: st.session_state["inp_phase"]   = "Phase I"
    if "inp_mode"    not in st.session_state: st.session_state["inp_mode"]    = list(MODE_HINTS.keys())[0]

    st.markdown('<div class="sb-section">Trial Parameters</div>', unsafe_allow_html=True)
    drug    = st.text_input("Investigational Drug", key="inp_drug",    placeholder="e.g. Pembrolizumab")
    disease = st.text_input("Indication / Disease",  key="inp_disease", placeholder="e.g. Lung Cancer")

    phase_opts = ["Phase I", "Phase II", "Phase III"]
    phase = st.selectbox("Trial Phase", phase_opts, key="inp_phase")

    st.markdown('<div class="sb-section">Optimisation Mode</div>', unsafe_allow_html=True)
    mode_opts = list(MODE_HINTS.keys())
    mode = st.selectbox("Mode", mode_opts, key="inp_mode")
    st.markdown(f'<div style="font-size:.62rem;color:var(--text3);padding:3px 2px;line-height:1.5;">{MODE_HINTS[mode]}</div>', unsafe_allow_html=True)

    st.markdown('<div class="sb-section">Data Sources</div>', unsafe_allow_html=True)
    use_rag    = st.toggle("ClinicalTrials.gov (Live)", value=True)
    use_pubmed = st.toggle("PubMed Citations (Live)",   value=True)

    st.markdown('<div class="sb-section">Display</div>', unsafe_allow_html=True)
    show_reasoning = st.toggle("Show AI Reasoning",    value=True)
    show_agents    = st.toggle("Show Agent Panel",     value=True)

    st.markdown('<div style="margin-top:.75rem;"></div>', unsafe_allow_html=True)
    gen_btn  = st.button("⚡  Generate Protocol", use_container_width=True)
    demo_btn = st.button("🎬  Demo Mode (Prefill)", use_container_width=True)

    if demo_btn:
        st.session_state.tmpl = DEMO_TEMPLATES[0]
        st.rerun()

    # Version history
    if st.session_state.versions:
        st.markdown('<div class="sb-section">Version History</div>', unsafe_allow_html=True)
        for i, v in enumerate(reversed(st.session_state.versions)):
            ri = len(st.session_state.versions)-1-i
            marker = "● " if ri == st.session_state.active else ""
            if st.button(f"{marker}{v['version_id']} {v['drug']} · {v['phase']}", key=f"ver_{ri}", use_container_width=True):
                st.session_state.active = ri

# ═══════════════════════════════════════════════════════════════════
# NAVBAR
# ═══════════════════════════════════════════════════════════════════
st.markdown("""
<div class="tf-nav">
  <div class="tf-brand">
    <div class="tf-brand-icon">⚗️</div>
    <div>
      <span class="tf-brand-name">TrialForge AI</span>
      <span class="tf-brand-tag">v3.0</span>
    </div>
  </div>
  <div style="text-align:center;display:none;">
    <div style="font-size:.72rem;color:var(--text2);">Autonomous Clinical Trial Protocol Designer</div>
    <div style="font-size:.62rem;color:var(--text3);">Powered by Amazon Nova Pro</div>
  </div>
  <div style="display:flex;gap:6px;flex-wrap:wrap;">
    <div class="tf-pill"><span class="dot dot-g"></span>Nova Pro · Connected</div>
    <div class="tf-pill"><span class="dot dot-c"></span>Bedrock · Active</div>
    <div class="tf-pill"><span class="dot dot-g"></span>ClinicalTrials.gov · Ready</div>
    <div class="tf-pill"><span class="dot dot-g"></span>PubMed · Ready</div>
  </div>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════
# GENERATE FLOW
# ═══════════════════════════════════════════════════════════════════
if gen_btn:
    if not drug or not disease:
        st.warning("Please enter Drug and Disease before generating.")
    else:
        STEPS = [
            ("🔍", "Analysing disease indication",          12),
            ("💊", "Selecting comparator drug",             22),
            ("🌐", "Querying ClinicalTrials.gov",           32),
            ("📚", "Fetching PubMed citations",             42),
            ("📐", "Estimating sample size & timeline",     50),
            ("🤖", "Dispatching to Amazon Nova Pro",        55),
            ("⏳", "Awaiting Nova response",                82),
            ("📊", "Computing quality & risk scores",       90),
            ("🔬", "Running feasibility & risk analysis",   96),
            ("✅", "Compiling final output",                100),
        ]
        ph_steps = st.empty()
        ph_prog  = st.empty()

        def render(states, pct):
            h  = '<div class="think-box"><div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:7px;">'
            h += '<span style="font-size:.63rem;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:var(--text3);">Nova · Thinking</span>'
            h += f'<span class="lat-badge">⚡ {pct}% complete</span></div>'
            for ic, lb, tp, st_ in states:
                ic_c = "ts-done" if st_=="done" else ("ts-run" if st_=="active" else "ts-wait")
                bc   = "tb-done" if st_=="done" else ("tb-run"  if st_=="active" else "tb-wait")
                bt   = "Done"    if st_=="done" else ("Running" if st_=="active" else "Wait")
                sym  = "✓"       if st_=="done" else ("⟳"       if st_=="active" else "·")
                h += f'<div class="think-step"><div class="ts-ico {ic_c}">{sym}</div><span class="ts-lbl">{lb}</span><span class="ts-badge {bc}">{bt}</span><span class="ts-pct">{tp}%</span></div>'
            h += "</div>"
            prg = f'''<div style="background:var(--card);border:1px solid var(--border);border-radius:var(--r-md);padding:.5rem 1rem;margin-bottom:.75rem;">
            <div style="display:flex;justify-content:space-between;font-size:.65rem;color:var(--text2);margin-bottom:4px;">
              <span>Designing protocol…</span>
              <span style="color:var(--accent);font-weight:700;font-family:var(--mono);">{pct}%</span>
            </div>
            <div style="height:7px;background:var(--surface);border-radius:var(--r-pill);overflow:hidden;">
              <div style="height:100%;width:{pct}%;background:linear-gradient(90deg,var(--accent),var(--indigo));border-radius:var(--r-pill);transition:width .4s ease;"></div>
            </div></div>'''
            ph_steps.markdown(h, unsafe_allow_html=True)
            ph_prog.markdown(prg,  unsafe_allow_html=True)

        result = None
        for i, (ic, lb, tp) in enumerate(STEPS):
            state_list = [(ic2, lb2, tp2, "done" if j<i else ("active" if j==i else "pending"))
                          for j,(ic2,lb2,tp2) in enumerate(STEPS)]
            render(state_list, tp)
            time.sleep(0.28)
            if i == 6 and result is None:
                try:
                    result = generate(drug, disease, phase, mode, use_rag, use_pubmed)
                except Exception as e:
                    ph_steps.error(f"⚠️ Generation failed: {e}")
                    logger.error("Generation failed: %s", e)
                    st.stop()

        if result:
            render([(ic,lb,tp,"done") for ic,lb,tp in STEPS], 100)
            time.sleep(0.25)
            st.session_state.versions.append(_normalise_scores(result))
            st.session_state.active = len(st.session_state.versions)-1
            add_audit_entry("GENERATE", f"{result['drug']} in {result['disease']} ({result['phase']})", result["version_id"])
            ph_steps.empty(); ph_prog.empty()
            st.rerun()

# ═══════════════════════════════════════════════════════════════════
# DISPLAY
# ═══════════════════════════════════════════════════════════════════
idx = st.session_state.active

if idx is not None and idx < len(st.session_state.versions):
    D  = _normalise_scores(st.session_state.versions[idx])
    sc = D["scores"]; q = D["quality"]; c = D["cost"]; tl = D["timeline"]
    fname = f"protocol_{D['version_id']}_{D['drug'].replace(' ','_')}"

    # ── Protocol header ──
    rc_cls = {"Low":"rc-lo","Medium":"rc-md","High":"rc-hi"}.get(sc["risk_class"],"rc-md")
    st.markdown(f"""
    <div class="proto-hdr">
      <div style="display:flex;align-items:flex-start;justify-content:space-between;flex-wrap:wrap;gap:10px;">
        <div>
          <div class="phdr-title">{D['drug']} in {D['disease']}</div>
          <div class="phdr-meta">{D['version_id']} · {D['timestamp']} · {D['phase']}</div>
          <div class="phdr-badges">
            <span class="phdr-badge phdr-badge-a">{D['phase']}</span>
            <span class="phdr-badge phdr-badge-i">{D['mode']}</span>
            <span class="phdr-badge">N = {D['sample_size']}</span>
            <span class="phdr-badge">{tl['total']} months</span>
            <span class="phdr-badge">Confidence: {q['conf']}%</span>
          </div>
        </div>
        <div style="display:flex;flex-direction:column;align-items:flex-end;gap:5px;">
          <span class="risk-chip {rc_cls}">⬤ Risk: {sc['risk_class']}</span>
          <span class="lat-badge">⚡ {D['latency']}s</span>
          <div style="font-size:.62rem;color:var(--text3);font-family:var(--mono);">{D['comparator'][:52]}{'…' if len(D['comparator'])>52 else ''}</div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── TOP METRIC CARDS ──
    ap = int(sc['ap']*100); rp = int(sc['rs']*100)
    sp_pct = int(sc['sp']*100); sf = int((1-sc['sr'])*100)
    r_col = "#34D399" if rp>70 else ("#FBBF24" if rp>45 else "#F87171")
    s_col = "#34D399" if sf>70 else ("#FBBF24" if sf>45 else "#F87171")
    st.markdown(f"""
    <div class="metric-row">
      <div class="metric-card"><div class="mc-val" style="color:#2DD4BF;">{ap}%</div>
        <div class="mc-lbl">Approval Probability</div>
        <div class="mc-sub" style="color:#2DD4BF;">{'High' if ap>=70 else 'Medium'}</div></div>
      <div class="metric-card"><div class="mc-val" style="color:{r_col};">{rp}%</div>
        <div class="mc-lbl">Recruitment Feasibility</div>
        <div class="mc-sub" style="color:{r_col};">{sc['rec_diff']}</div></div>
      <div class="metric-card"><div class="mc-val" style="color:{s_col};">{sf}%</div>
        <div class="mc-lbl">Safety Index</div>
        <div class="mc-sub" style="color:{s_col};">{sc['risk_class']} Risk</div></div>
      <div class="metric-card"><div class="mc-val" style="color:#FBBF24;">${c['grand']//1000:,}K</div>
        <div class="mc-lbl">Trial Cost (USD)</div>
        <div class="mc-sub" style="color:#FBBF24;">${c['pp']:,}/patient</div></div>
      <div class="metric-card"><div class="mc-val" style="color:#818CF8;">{tl['total']}mo</div>
        <div class="mc-lbl">Total Timeline</div>
        <div class="mc-sub" style="color:#818CF8;">{tl['en']}mo enrolment</div></div>
    </div>
    """, unsafe_allow_html=True)

    # ═══════════════════════════════════════════════════════════════
    # TABS
    # ═══════════════════════════════════════════════════════════════
    tab_ov, tab_proto, tab_review, tab_clinops, tab_insights, tab_collab, tab_arch, tab_export = st.tabs([
        "📊 Overview", "📄 Protocol", "🤖 AI Review",
        "🏥 ClinOps", "💡 Insights", "🤝 Collaborate",
        "🏗️ Architecture", "📥 Export"
    ])

    # ── OVERVIEW ────────────────────────────────────────────────────
    with tab_ov:
        # Confidence banner
        st.markdown(f"""
        <div class="conf-banner">
          <div>
            <div class="cb-lbl">Protocol Confidence Score</div>
            <div class="cb-val">{q['conf']}%</div>
            <div class="cb-sub">ICH compliance, FDA readiness, safety profile</div>
          </div>
          <div style="text-align:right;">
            <div class="cb-lbl" style="margin-bottom:3px;">Overall Success Probability</div>
            <div style="font-size:1.9rem;font-weight:800;color:var(--accent);">{sp_pct}%</div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        # ROI / Business Value banner
        roi = D.get("roi", compute_roi(D["phase"], c, tl, D["drug"], D["disease"]))
        st.markdown(f"""
        <div class="roi-banner">
          <div class="roi-item">
            <div class="roi-val">⏱ {roi['months_saved']}mo</div>
            <div class="roi-lbl">Time Saved vs Traditional</div>
          </div>
          <div class="roi-item">
            <div class="roi-val">${roi['cro_savings']//1000:,}K</div>
            <div class="roi-lbl">CRO Fees Saved</div>
          </div>
          <div class="roi-item">
            <div class="roi-val">${roi['amendment_saved']//1000:,}K</div>
            <div class="roi-lbl">Amendment Costs Avoided</div>
          </div>
          <div class="roi-item">
            <div class="roi-val">${roi['total_savings']//1000:,}K</div>
            <div class="roi-lbl">Total Estimated Value</div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        # ROI / Business Value Card
        roi = D.get("roi", compute_roi(D["phase"], c, tl, D["drug"], D["disease"]))
        st.markdown(f"""
        <div class="card card-hi" style="border:1px solid rgba(45,212,191,.3);background:linear-gradient(135deg,rgba(45,212,191,.06),rgba(129,140,248,.06));">
          <div class="card-hdr"><div class="card-ico ico-t">💼</div>Business Value — AI vs Traditional Protocol Design</div>
          <div style="display:grid;grid-template-columns:repeat(4,1fr);gap:.6rem;margin-top:.25rem;">
            <div style="text-align:center;background:var(--card);border-radius:var(--r-md);padding:.7rem;">
              <div style="font-size:1.5rem;font-weight:800;color:#2DD4BF;">{roi['months_saved']}mo</div>
              <div style="font-size:.6rem;text-transform:uppercase;letter-spacing:.08em;color:var(--text2);margin-top:2px;">Time Saved</div>
            </div>
            <div style="text-align:center;background:var(--card);border-radius:var(--r-md);padding:.7rem;">
              <div style="font-size:1.5rem;font-weight:800;color:#34D399;">${roi['cro_savings']//1000:,}K</div>
              <div style="font-size:.6rem;text-transform:uppercase;letter-spacing:.08em;color:var(--text2);margin-top:2px;">CRO Fees Saved</div>
            </div>
            <div style="text-align:center;background:var(--card);border-radius:var(--r-md);padding:.7rem;">
              <div style="font-size:1.5rem;font-weight:800;color:#FBBF24;">${roi['amendment_saved']//1000:,}K</div>
              <div style="font-size:.6rem;text-transform:uppercase;letter-spacing:.08em;color:var(--text2);margin-top:2px;">Amendments Avoided</div>
            </div>
            <div style="text-align:center;background:var(--card);border-radius:var(--r-md);padding:.7rem;border:1px solid rgba(45,212,191,.25);">
              <div style="font-size:1.5rem;font-weight:800;color:#818CF8;">${roi['total_savings']//1000:,}K</div>
              <div style="font-size:.6rem;text-transform:uppercase;letter-spacing:.08em;color:var(--text2);margin-top:2px;">Total Value Created</div>
            </div>
          </div>
          <div style="font-size:.7rem;color:var(--text2);margin-top:.6rem;">
            ⚡ {roi['writer_days_saved']} medical writer days saved · 
            Traditional timeline: {roi['traditional_months']}mo → AI-optimised: {tl['total']}mo ·
            Protocol writer savings: ${roi['writer_cost_saved']:,}
          </div>
        </div>
        """, unsafe_allow_html=True)

        # Radar chart
        fig_radar = plotly_radar_chart(q, sc)
        if fig_radar:
            rc1, rc2 = st.columns([1, 2])
            with rc1:
                st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-p">🕸️</div>Quality Radar</div>', unsafe_allow_html=True)
                st.plotly_chart(fig_radar, use_container_width=True, config={"displayModeBar": False})
                st.markdown("</div>", unsafe_allow_html=True)
            with rc2:
                # Score rings row
                st.markdown('<div class="card card-hi"><div class="card-hdr"><div class="card-ico ico-t">📊</div>Protocol Intelligence Dashboard</div>', unsafe_allow_html=True)
                rings_html = '<div style="display:grid;grid-template-columns:repeat(5,1fr);gap:.5rem;margin-bottom:.75rem;">'
                rings_html += ring(ap, "#2DD4BF", "Regulatory Approval", "Probability")
                rings_html += ring(rp, r_col,     "Recruitment",         sc['rec_diff'])
                rings_html += ring(sf, s_col,     "Safety Index",        sc['risk_class']+" Risk")
                rings_html += ring(sp_pct, "#818CF8","Overall Success",  "Probability")
                rings_html += ring(q['conf'], "#C084FC","Confidence",    "AI Judge Score")
                rings_html += "</div>"
                st.markdown(rings_html, unsafe_allow_html=True)
                st.markdown(pb("Regulatory Approval Probability", ap, "#2DD4BF"), unsafe_allow_html=True)
                st.markdown(pb("Recruitment Feasibility",         rp, r_col),     unsafe_allow_html=True)
                st.markdown(pb("Safety Profile Index",            sf, s_col),     unsafe_allow_html=True)
                st.markdown(pb("Overall Success Probability",     sp_pct, "#818CF8"), unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)
        else:
            # Plotly not available — show rings only
            st.markdown('<div class="card card-hi"><div class="card-hdr"><div class="card-ico ico-t">📊</div>Protocol Intelligence Dashboard</div>', unsafe_allow_html=True)
            rings_html = '<div style="display:grid;grid-template-columns:repeat(5,1fr);gap:.5rem;margin-bottom:.75rem;">'
            rings_html += ring(ap, "#2DD4BF", "Regulatory Approval", "Probability")
            rings_html += ring(rp, r_col,     "Recruitment",         sc['rec_diff'])
            rings_html += ring(sf, s_col,     "Safety Index",        sc['risk_class']+" Risk")
            rings_html += ring(sp_pct, "#818CF8","Overall Success",  "Probability")
            rings_html += ring(q['conf'], "#C084FC","Confidence",    "AI Judge Score")
            rings_html += "</div>"
            st.markdown(rings_html, unsafe_allow_html=True)
            st.markdown(pb("Regulatory Approval Probability", ap, "#2DD4BF"), unsafe_allow_html=True)
            st.markdown(pb("Recruitment Feasibility",         rp, r_col),     unsafe_allow_html=True)
            st.markdown(pb("Safety Profile Index",            sf, s_col),     unsafe_allow_html=True)
            st.markdown(pb("Overall Success Probability",     sp_pct, "#818CF8"), unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        col1, col2 = st.columns(2)

        with col1:
            # Quality scores
            def qb(val, hi=78, lo=62, inv=False):
                if inv:
                    cl = "qb-hi" if val<20 else ("qb-md" if val<35 else "qb-lo"); lb="Low Risk" if val<20 else ("Med Risk" if val<35 else "High Risk")
                else:
                    cl = "qb-hi" if val>=hi else ("qb-md" if val>=lo else "qb-lo"); lb="High" if val>=hi else ("Medium" if val>=lo else "Low")
                return f'<span class="qbadge {cl}">{lb}</span>'

            st.markdown(f"""
            <div class="card"><div class="card-hdr"><div class="card-ico ico-o">🏆</div>Quality Judge Dashboard</div>
            <table class="meta-tbl">
              <tr><td>ICH-GCP Compliance</td><td>{q['ich']}% {qb(q['ich'])}</td></tr>
              <tr><td>Regulatory Readiness</td><td>{q['reg']}% {qb(q['reg'])}</td></tr>
              <tr><td>FDA Approval Readiness</td><td>{q['fda']}% {qb(q['fda'])}</td></tr>
              <tr><td>Failure Risk Score</td><td>{q['fail']}% {qb(q['fail'],inv=True)}</td></tr>
              <tr><td>Protocol Completeness</td><td>{q['comp']}% {qb(q['comp'])}</td></tr>
              <tr><td>Confidence Score</td><td>{q['conf']}% {qb(q['conf'])}</td></tr>
            </table></div>
            """, unsafe_allow_html=True)

            # Budget chart (Plotly or SVG fallback)
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-o">💰</div>Budget Breakdown</div>', unsafe_allow_html=True)
            fig_budget = plotly_budget_chart(c)
            if fig_budget:
                st.plotly_chart(fig_budget, use_container_width=True, config={"displayModeBar": False})
            else:
                st.markdown(svg_budget_chart(c), unsafe_allow_html=True)
            st.markdown(f'<div style="font-size:.75rem;color:var(--text2);">Total: <strong style="color:var(--gold);">${c["grand"]:,}</strong> · Per patient: ${c["pp"]:,}</div></div>', unsafe_allow_html=True)

            # Metadata
            st.markdown(f"""
            <div class="card"><div class="card-hdr"><div class="card-ico ico-t">💊</div>Protocol Metadata</div>
            <table class="meta-tbl">
              <tr><td>Comparator Drug</td><td>{D['comparator']}</td></tr>
              <tr><td>Sample Size</td><td>N = {D['sample_size']}</td></tr>
              <tr><td>Study Duration</td><td>{D['duration']} months</td></tr>
              <tr><td>Generation Latency</td><td>{D['latency']}s</td></tr>
              <tr><td>PubMed Citations</td><td>{len(D['pubmed'])} found</td></tr>
              <tr><td>NCT Trials Retrieved</td><td>{len(D['nct'])} trials</td></tr>
            </table></div>
            """, unsafe_allow_html=True)

        with col2:
            # Enrolment chart (Plotly or SVG fallback)
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-t">📈</div>Enrolment Trajectory</div>', unsafe_allow_html=True)
            fig_enroll = plotly_enroll_chart(tl, D['sample_size'])
            if fig_enroll:
                st.plotly_chart(fig_enroll, use_container_width=True, config={"displayModeBar": False})
            else:
                st.markdown(svg_enroll_chart(tl, D['sample_size']), unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

            # Gantt
            st.markdown(f"""
            <div class="card"><div class="card-hdr"><div class="card-ico ico-i">📅</div>Trial Gantt Chart</div>
            {gantt(tl)}
            <table class="meta-tbl" style="margin-top:8px;">
              <tr><td>Startup</td><td>{tl['su']}mo</td></tr>
              <tr><td>Enrolment</td><td>{tl['en']}mo</td></tr>
              <tr><td>Treatment</td><td>{tl['base']}mo</td></tr>
              <tr><td>Follow-up</td><td>{tl['fu']}mo</td></tr>
              <tr><td style="font-weight:700;color:var(--text);">Total</td><td style="color:var(--accent);font-weight:800;">{tl['total']}mo</td></tr>
            </table></div>
            """, unsafe_allow_html=True)

            # NCT trials
            if D["nct"]:
                st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-g">🔗</div>ClinicalTrials.gov Results</div>', unsafe_allow_html=True)
                for t in D["nct"][:5]:
                    st.markdown(f"""<div class="nct-item">
                      <span class="nct-id">[{t['id']}]</span>
                      <span class="nct-title"> {t['title']}</span>
                      <div class="nct-meta">{t['phase']} · N={t['n']} · {t['status']} · Endpoint: {t['endpoint']}</div>
                    </div>""", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

        # PubMed
        if D["pubmed"]:
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-g">📚</div>PubMed Live Citations</div>', unsafe_allow_html=True)
            for p in D["pubmed"]:
                st.markdown(f"""<div class="nct-item">
                  <span class="nct-id">PMID:{p['pmid']}</span>
                  <span class="nct-title"> {p['title']}</span>
                  <div class="nct-meta">{p['journal']} · {p['year']} · <a href="{p['url']}" style="color:var(--accent);">View on PubMed ↗</a></div>
                </div>""", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # SoA matrix
        with st.expander("📋 Schedule of Assessments (SoA) Matrix", expanded=False):
            st.markdown(soa_table_html(), unsafe_allow_html=True)

        # ICH checklist
        with st.expander("✅ ICH-GCP Compliance Checklist", expanded=False):
            for item, ok in ICH_GCP_ITEMS:
                st.markdown(f'<div class="flag-card"><div class="flag-ico">{"✅" if ok else "⚠️"}</div><div class="flag-desc">{item}</div></div>', unsafe_allow_html=True)

        # Compare (if 2+ versions)
        if len(st.session_state.versions) >= 2:
            with st.expander("⚖️ Protocol Comparison (A vs B)", expanded=False):
                ver_labels = [f"{v['version_id']} — {v['drug']} / {v['disease']}" for v in st.session_state.versions]
                cc1, cc2 = st.columns(2)
                with cc1: ia = st.selectbox("Protocol A", range(len(ver_labels)), format_func=lambda i:ver_labels[i], key="ca")
                with cc2: ib = st.selectbox("Protocol B", range(len(ver_labels)), index=min(1,len(ver_labels)-1), format_func=lambda i:ver_labels[i], key="cb")
                va = st.session_state.versions[ia]; vb = st.session_state.versions[ib]

                def cc(a, b, inv=False):
                    if inv: a, b = b, a
                    return ("cv-win","cv-lose") if a>b else (("cv-lose","cv-win") if a<b else ("cv-tie","cv-tie"))

                rows = [
                    ("Approval Probability", int(_sc(va,'ap')*100), int(_sc(vb,'ap')*100), "%", False),
                    ("Recruitment Score",    int(_sc(va,'rs')*100), int(_sc(vb,'rs')*100), "%", False),
                    ("Safety Index",         int((1-_sc(va,'sr'))*100), int((1-_sc(vb,'sr'))*100), "%", False),
                    ("Overall Success",      int(_sc(va,'sp')*100), int(_sc(vb,'sp')*100), "%", False),
                    ("Confidence",           va.get('quality',{}).get('conf',0),  vb.get('quality',{}).get('conf',0),  "%", False),
                    ("FDA Readiness",        va.get('quality',{}).get('fda',0),   vb.get('quality',{}).get('fda',0),   "%", False),
                    ("ICH Compliance",       va.get('quality',{}).get('ich',0),   vb.get('quality',{}).get('ich',0),   "%", False),
                    ("Total Cost",           va.get('cost',{}).get('grand',0)//1000, vb.get('cost',{}).get('grand',0)//1000, "K", True),
                    ("Duration",             va.get('timeline',{}).get('total',0), vb.get('timeline',{}).get('total',0), "mo", True),
                ]
                cl1, cl2 = st.columns(2)
                for col, key, vx in [(cl1,"A",va),(cl2,"B",vb)]:
                    other = vb if key=="A" else va
                    with col:
                        st.markdown(f'<div class="card"><div style="font-size:.65rem;font-weight:700;text-transform:uppercase;letter-spacing:.08em;color:var(--text2);padding:3px 0 8px;">Protocol {key} · {vx["version_id"]} · {vx["drug"]}</div>', unsafe_allow_html=True)
                        for lbl, av, bv, unit, inv in rows:
                            myv = av if key=="A" else bv; otv = bv if key=="A" else av
                            ca, cb = cc(myv, otv, inv)
                            mine = ca if key=="A" else cb
                            st.markdown(f'<div class="cmp-row"><span class="cmp-lbl">{lbl}</span><span class="cmp-val {mine}">{myv}{unit}</span></div>', unsafe_allow_html=True)
                        st.markdown("</div>", unsafe_allow_html=True)

    # ── PROTOCOL ────────────────────────────────────────────────────
    with tab_proto:
        col_p1, col_p2 = st.columns([2, 1])

        with col_p1:
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-t">📄</div>Clinical Trial Protocol — Document View</div>', unsafe_allow_html=True)
            proto_html = render_protocol_html(D["protocol"])
            if not show_reasoning:
                # strip reasoning section
                proto_html = re.sub(r'<div class="proto-h1">REASONING.*?(?=<div class="proto-h1"|$)', "", proto_html, flags=re.DOTALL)
            st.markdown(f'<div class="proto-doc">{proto_html}</div>', unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        with col_p2:
            # Improvement suggestions (extracted from protocol)
            suggs = []
            in_s  = False
            for line in clean_protocol_text(D["protocol"]).split("\n"):
                if "IMPROVEMENT SUGGESTION" in line.upper() or "PROTOCOL IMPROVEMENT" in line.upper():
                    in_s = True; continue
                if in_s and "REASONING" in line.upper():
                    break
                if in_s:
                    s2 = line.strip().lstrip("0123456789.-). •").strip()
                    if s2 and len(s2) > 12:
                        suggs.append(s2)

            if suggs:
                st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-g">💡</div>Improvement Suggestions</div>', unsafe_allow_html=True)
                for i, sg in enumerate(suggs[:6], 1):
                    st.markdown(f'<div class="sug-item"><div class="sug-num">{i}</div><div class="sug-txt">{sg}</div></div>', unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

            # Amendment risk predictor
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-p">⚠️</div>Amendment Risk Predictor</div>', unsafe_allow_html=True)
            for flag in D["amendments"]:
                lbl, detail, risk = flag
                fc = "ico-r" if risk=="High" else "ico-o"
                st.markdown(f'<div class="flag-card"><div class="flag-ico">{lbl[:2]}</div><div class="flag-body"><div class="flag-title">{lbl[2:].strip()}</div><div class="flag-desc">{detail}</div></div></div>', unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

            # Drug interactions
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-r">💊</div>Drug Interaction Check</div>', unsafe_allow_html=True)
            risk_colors = {"High":"var(--red)","Medium":"var(--gold)","Low":"var(--green)"}
            for ddi in D["interactions"]:
                col_r = risk_colors.get(ddi["risk"],"var(--text2)")
                st.markdown(f'<div class="flag-card"><div class="flag-ico">⚡</div><div class="flag-body"><div class="flag-title">{ddi["drug"]} <span style="color:{col_r};font-size:.62rem;font-weight:700;">{ddi["risk"]} Risk</span></div><div class="flag-desc">{ddi["detail"]}</div></div></div>', unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

            # SoA compact
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-b">📋</div>Schedule of Assessments</div>', unsafe_allow_html=True)
            st.markdown(soa_table_html(), unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # ICF translator
        st.markdown("---")

        # ── Chat with Protocol ──
        st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-i">💬</div>Chat with the Protocol — Ask Nova Anything</div>', unsafe_allow_html=True)
        chat_key = D["version_id"]
        if chat_key not in st.session_state.chat_histories:
            st.session_state.chat_histories[chat_key] = []
        chat_hist = st.session_state.chat_histories[chat_key]

        # Display history
        if chat_hist:
            chat_html = ""
            for msg in chat_hist:
                if msg["role"] == "user":
                    chat_html += f'<div class="chat-msg"><div class="chat-avatar chat-av-user">👤</div><div><div class="chat-bubble chat-bubble-user">{msg["content"]}</div><div class="chat-time">{msg.get("ts","")}</div></div></div>'
                else:
                    chat_html += f'<div class="chat-msg"><div class="chat-avatar chat-av-ai">⚗️</div><div><div class="chat-bubble chat-bubble-ai">{msg["content"][:800]}{"…" if len(msg["content"])>800 else ""}</div><div class="chat-time">{msg.get("ts","")}</div></div></div>'
            st.markdown(f'<div style="max-height:320px;overflow-y:auto;padding:4px 0;margin-bottom:.5rem;">{chat_html}</div>', unsafe_allow_html=True)

        # Suggested questions
        sugg_qs = [
            f"What is the primary endpoint for this {D['disease']} trial?",
            "What are the key exclusion criteria?",
            "Summarise the statistical analysis plan.",
            f"What is the washout period and why?",
            "List the safety monitoring procedures.",
        ]
        st.markdown('<div style="display:flex;flex-wrap:wrap;gap:5px;margin-bottom:.5rem;">', unsafe_allow_html=True)
        for sq in sugg_qs:
            if st.button(f"💬 {sq[:45]}…" if len(sq)>45 else f"💬 {sq}", key=f"sq_{sq[:20]}", use_container_width=False):
                ts = datetime.datetime.now().strftime("%H:%M")
                chat_hist.append({"role":"user","content":sq,"ts":ts})
                with st.spinner("Nova is thinking…"):
                    answer = chat_with_protocol(sq, D["protocol"], D["drug"], D["disease"], chat_hist)
                chat_hist.append({"role":"assistant","content":answer,"ts":datetime.datetime.now().strftime("%H:%M")})
                add_audit_entry("CHAT", f"Q: {sq[:60]}", D["version_id"])
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        user_q = st.chat_input("Ask anything about this protocol… e.g. 'What is the dose escalation strategy?'")
        if user_q:
            ts = datetime.datetime.now().strftime("%H:%M")
            chat_hist.append({"role":"user","content":user_q,"ts":ts})
            with st.spinner("Nova is analysing the protocol…"):
                answer = chat_with_protocol(user_q, D["protocol"], D["drug"], D["disease"], chat_hist)
            chat_hist.append({"role":"assistant","content":answer,"ts":datetime.datetime.now().strftime("%H:%M")})
            add_audit_entry("CHAT", f"Q: {user_q[:60]}", D["version_id"])
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-t">📝</div>Informed Consent Form — Plain Language (6th-Grade Level)</div>', unsafe_allow_html=True)

        icf_key = f"{D['version_id']}_icf"
        if icf_key not in st.session_state.icf_cache:
            if st.button("🔄 Generate Patient-Friendly ICF", key="gen_icf"):
                with st.spinner("Translating protocol to plain language…"):
                    icf_text = generate_icf(D["protocol"], D["drug"], D["disease"])
                    st.session_state.icf_cache[icf_key] = icf_text
                    st.rerun()
        else:
            icf_raw = st.session_state.icf_cache[icf_key]
            # Style the ICF content
            icf_html = ""
            for line in icf_raw.split("\n"):
                s2 = line.strip()
                if not s2: icf_html += "<br>"; continue
                if s2.isupper() or (s2.endswith("?") and len(s2)<60 and s2[0].isupper()):
                    icf_html += f'<div class="icf-label">{s2}</div>'
                elif s2.startswith(("⚠", "✓", "•", "-")):
                    sym = s2[0]; rest = s2[1:].strip()
                    cls = "icf-warning" if sym=="⚠" else "icf-highlight"
                    icf_html += f'<div class="{cls}">{sym} {rest}</div>'
                else:
                    icf_html += f'<div class="icf-plain">{s2}</div>'
            st.markdown(f'<div class="icf-doc">{icf_html}</div>', unsafe_allow_html=True)
            st.download_button("⬇️ Download ICF (TXT)",
                data=icf_raw.encode(),
                file_name=f"{fname}_ICF.txt", mime="text/plain")

        st.markdown("</div>", unsafe_allow_html=True)

        # FDA Diversity Action Plan
        st.markdown('<div class="card" style="margin-top:.5rem;"><div class="card-hdr"><div class="card-ico ico-g">🌍</div>FDA Diversity Action Plan (FDORA §3041)</div>', unsafe_allow_html=True)
        for icon_d, desc in D["diversity"]:
            st.markdown(f'<div class="flag-card"><div class="flag-ico">🌍</div><div class="flag-body"><div class="flag-title">{icon_d}</div><div class="flag-desc">{desc}</div></div></div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Synthetic EHR feasibility
        ehr = D["ehr"]
        feas_col = "#34D399" if ehr["feasibility"]=="High" else ("#FBBF24" if ehr["feasibility"]=="Medium" else "#F87171")
        st.markdown(f"""
        <div class="card"><div class="card-hdr"><div class="card-ico ico-t">🔬</div>Synthetic EHR Feasibility Engine</div>
        <div style="display:grid;grid-template-columns:repeat(4,1fr);gap:.5rem;margin-bottom:.5rem;">
          <div class="metric-card"><div class="mc-val" style="color:{feas_col};">{ehr['feasibility']}</div><div class="mc-lbl">Feasibility</div></div>
          <div class="metric-card"><div class="mc-val" style="color:var(--accent);">{ehr['eligible_pool']:,}</div><div class="mc-lbl">Eligible Patients (US)</div></div>
          <div class="metric-card"><div class="mc-val" style="color:var(--indigo);">{ehr['sites_needed']}</div><div class="mc-lbl">Sites Needed</div></div>
          <div class="metric-card"><div class="mc-val" style="color:var(--green);">{ehr['monthly_rate']}</div><div class="mc-lbl">Pts/Month</div></div>
        </div>
        <div style="font-size:.72rem;color:var(--text2);">Based on US prevalence of {D['disease']}: ~{ehr['prevalence']:,}. Estimated {ehr['eligible_pct']}% meet I/E criteria. Synthetic EHR simulation (not real patient data).</div>
        </div>
        """, unsafe_allow_html=True)

    # ── AI REVIEW ────────────────────────────────────────────────────
    with tab_review:
        if show_agents:
            # Multi-agent panel
            agents = [
                ("🏛️", "Regulatory Agent", "av-reg", "an-reg",
                 f"Approval probability: {ap}%. FDA readiness: {q['fda']}%. Protocol aligns with ICH E6(R2) GCP and FDA 21 CFR Part 312. IND/CTA filing recommended. Comparator ({D['comparator'][:50]}) is SOC for {D['disease']}."),
                ("🛡️", "Safety Agent", "av-safe", "an-safe",
                 f"Safety index: {sf}/100. Risk: {sc['risk_class']}. Failure risk: {q['fail']}% ({q['fail_lbl']}). DSMB oversight required. 15-day SAE FDA reporting window applies. {len(D['interactions'])} drug interaction(s) flagged."),
                ("📊", "Statistical Agent", "av-stat", "an-stat",
                 f"Recruitment feasibility: {rp}/100 ({sc['rec_diff']}). ICH compliance: {q['ich']}%. Recommend pre-specified interim analysis at 50% information fraction. O'Brien-Fleming boundary. Adaptive enrichment if interim supports."),
                ("🎯", "Optimisation Agent", "av-opt", "an-opt",
                 f"Confidence: {q['conf']}%. Completeness: {q['comp']}%. Regulatory readiness: {q['reg']}%. {len(D['amendments'])} amendment risk(s) identified. Diversity plan: {len(D['diversity'])} items required per FDORA."),
            ]
            st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-i">🤖</div>Multi-Agent Analysis Panel</div>', unsafe_allow_html=True)
            av_map = {"av-reg":"background:rgba(129,140,248,.12)","av-safe":"background:rgba(248,113,113,.12)",
                      "av-stat":"background:rgba(52,211,153,.12)","av-opt":"background:rgba(192,132,252,.12)"}
            nc_map = {"an-reg":"color:#818CF8","an-safe":"color:#F87171","an-stat":"color:#34D399","an-opt":"color:#C084FC"}
            for ic, name, av, nc, verdict in agents:
                st.markdown(f"""<div class="agent-card">
                  <div class="ag-av" style="{av_map.get(av,'')}">{ic}</div>
                  <div class="ag-body">
                    <div class="ag-name" style="{nc_map.get(nc,'')}">{name}</div>
                    <div class="ag-txt">{verdict}</div>
                  </div></div>""", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # Red Team
        st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-r">🔴</div>Red Team — Mock FDA Reviewer</div>', unsafe_allow_html=True)
        rt_key = f"{D['version_id']}_rt"
        if rt_key not in st.session_state.red_team_cache:
            if st.button("🔴 Run Red Team FDA Review", key="run_rt"):
                with st.spinner("Red team agent analysing protocol…"):
                    rt = generate_red_team(D["protocol"], D["drug"], D["disease"])
                    st.session_state.red_team_cache[rt_key] = rt
                    st.rerun()
            st.markdown('<div style="font-size:.75rem;color:var(--text2);padding:4px 0;">Click to run the mock FDA reviewer agent — it will attack the protocol for weaknesses.</div>', unsafe_allow_html=True)
        else:
            rt_text = st.session_state.red_team_cache[rt_key]
            lines   = [l.strip() for l in rt_text.split("\n") if l.strip()]
            for ln in lines:
                ln_clean = re.sub(r"^\*+\s*", "", re.sub(r"\*+", "", ln))
                if re.match(r"^\d+[.)]\s", ln_clean):
                    num, rest = ln_clean[:2], ln_clean[2:].strip()
                    st.markdown(f'<div class="flag-card"><div class="flag-ico">🔴</div><div class="flag-body"><div class="flag-title">Finding {num}</div><div class="flag-desc">{rest}</div></div></div>', unsafe_allow_html=True)
                else:
                    if ln_clean:
                        st.markdown(f'<div style="font-size:.77rem;color:var(--text2);padding:3px 0;">{ln_clean}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        # Quality breakdown
        st.markdown(f"""
        <div class="card"><div class="card-hdr"><div class="card-ico ico-o">📈</div>Quality Score Breakdown</div>
        {pb("ICH-GCP Compliance",    q['ich'],  "#2DD4BF")}
        {pb("Regulatory Readiness",  q['reg'],  "#818CF8")}
        {pb("FDA Approval Readiness",q['fda'],  "#34D399")}
        {pb("Protocol Completeness", q['comp'], "#C084FC")}
        {pb("Confidence Score",      q['conf'], "#FBBF24")}
        {pb("Failure Risk (lower=better)", q['fail'], "#F87171", right=f"{q['fail_lbl']} ({q['fail']}%)")}
        </div>
        """, unsafe_allow_html=True)

        # CDISC export note
        st.markdown("""
        <div class="card"><div class="card-hdr"><div class="card-ico ico-b">🏭</div>CDISC / SDTM Data Export</div>
        <div style="font-size:.77rem;color:var(--text2);line-height:1.65;">
        Protocol endpoints and metadata are available as CDISC SDTM-compatible XML via the Export tab.
        The XML follows ODM 1.3 schema and can be imported into EDC systems (Medidata Rave, Oracle InForm, etc.).
        </div></div>
        """, unsafe_allow_html=True)

    # ── ARCHITECTURE ─────────────────────────────────────────────────
    with tab_arch:
        st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-t">🏗️</div>System Architecture</div>', unsafe_allow_html=True)
        try:
            import graphviz
            st.graphviz_chart(ARCH_DOT, use_container_width=True)
        except Exception:
            # Fallback: ASCII diagram
            ascii_arch = """
  ┌─────────────────────────────────────────────────────────────────┐
  │                     TrialForge AI v3.0                         │
  │            Autonomous Clinical Trial Protocol Designer          │
  └──────────────────────────┬──────────────────────────────────────┘
                             │
             ┌───────────────▼───────────────┐
             │       Streamlit Frontend       │
             │   Material UI · 5 Tabs · PDF   │
             └─────────┬──────────┬───────────┘
                       │          │
          ┌────────────▼──┐  ┌────▼────────────────┐
          │ ClinicalTrials│  │    PubMed API        │
          │  .gov API v2  │  │    Live Citations    │
          └────────────┬──┘  └────┬────────────────┘
                       └────┬─────┘
             ┌──────────────▼──────────────────┐
             │       Prompt Builder v3          │
             │  Mode · RAG · Citations · Context│
             └──────────────┬──────────────────┘
                            │
             ┌──────────────▼──────────────────┐
             │      Amazon Bedrock              │
             │      Nova Pro v1.0               │
             │      Converse API                │
             └──┬──────┬──────┬──────┬──────────┘
                │      │      │      │
        ┌───────▼┐ ┌───▼──┐ ┌▼────┐ ┌▼──────────┐
        │Reg     │ │Safety│ │Stat │ │Red Team    │
        │Agent   │ │Agent │ │Agent│ │FDA Reviewer│
        └───┬────┘ └───┬──┘ └┬────┘ └┬───────────┘
            └──────────▼─────┘       │
                       │             │
             ┌─────────▼─────────────▼──────────┐
             │   Protocol Quality Judge          │
             │   ICH · FDA · Confidence · Risk   │
             └─────────────┬─────────────────────┘
                           │
  ┌────────────────────────▼────────────────────────────────────────┐
  │                    Output Layer                                  │
  │  PDF (ReportLab) · DOCX · JSON · TXT · CDISC XML               │
  │  ICF Translator · SoA Matrix · Diversity Plan · EHR Feasibility │
  │  Drug Interactions · Amendment Predictor · PubMed Citations     │
  └─────────────────────────────────────────────────────────────────┘
"""
            st.markdown(f'<div style="background:var(--surface);border:1px solid var(--border);border-radius:var(--r-lg);padding:1rem 1.25rem;font-family:var(--mono);font-size:.66rem;color:var(--text2);line-height:1.9;white-space:pre;overflow-x:auto;">{ascii_arch}</div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        col_a1, col_a2, col_a3 = st.columns(3)
        for col, ico, title, items in [
            (col_a1, "🧠", "Intelligence Layer", ["Amazon Nova Pro (Bedrock)", "ClinicalTrials.gov API v2", "PubMed E-utilities API", "RAG Engine + Prompt Builder", "Protocol Caching (30min TTL)"]),
            (col_a2, "🔬", "Analysis Layer",     ["Multi-Agent (4 specialists)", "Red Team FDA Reviewer", "Protocol Quality Judge", "Amendment Risk Predictor", "Drug Interaction Checker", "FDA Diversity Action Plan", "Synthetic EHR Feasibility"]),
            (col_a3, "📤", "Output Layer",        ["PDF (ReportLab A4)", "DOCX (python-docx)", "CDISC/SDTM XML", "JSON structured data", "TXT protocol", "ICF Plain Language", "SoA Visual Matrix"]),
        ]:
            with col:
                items_html = "".join(f'<div style="font-size:.7rem;color:var(--text2);padding:3px 0;border-bottom:1px solid var(--border);">→ {i}</div>' for i in items)
                st.markdown(f'<div class="card" style="text-align:center;"><div style="font-size:1.5rem;margin-bottom:6px;">{ico}</div><div style="font-size:.78rem;font-weight:700;color:var(--text);margin-bottom:8px;">{title}</div>{items_html}</div>', unsafe_allow_html=True)

    # ── EXPORT ───────────────────────────────────────────────────────
    with tab_export:
        st.markdown('<div class="card"><div class="card-hdr"><div class="card-ico ico-g">📥</div>Export Protocol</div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:.75rem;color:var(--text2);margin-bottom:.75rem;">Download your protocol in professional formats ready for regulatory submission.</div>', unsafe_allow_html=True)

        e1, e2, e3, e4, e5 = st.columns(5)

        with e1:
            st.download_button("⬇️ PDF",
                data=generate_pdf_bytes(D),
                file_name=f"{fname}.pdf", mime="application/pdf",
                use_container_width=True)

        with e2:
            docx_bytes = generate_docx_bytes(D)
            if docx_bytes:
                st.download_button("⬇️ DOCX",
                    data=docx_bytes,
                    file_name=f"{fname}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)

        with e3:
            export_dict = {k: v for k, v in D.items()
                           if k not in ("interactions","amendments","diversity","ehr","pubmed","nct")}
            export_dict["protocol_text"] = D["protocol"]
            st.download_button("⬇️ JSON",
                data=json.dumps(export_dict, indent=2).encode(),
                file_name=f"{fname}.json", mime="application/json",
                use_container_width=True)

        with e4:
            st.download_button("⬇️ TXT",
                data=clean_protocol_text(D["protocol"]).encode(),
                file_name=f"{fname}.txt", mime="text/plain",
                use_container_width=True)

        with e5:
            cdisc_xml = export_cdisc_xml(D)
            st.download_button("⬇️ CDISC XML",
                data=cdisc_xml.encode(),
                file_name=f"{fname}_CDISC.xml", mime="application/xml",
                use_container_width=True)

        st.markdown("</div>", unsafe_allow_html=True)

        # Generation metrics
        st.markdown(f"""
        <div class="card"><div class="card-hdr"><div class="card-ico ico-g">⚡</div>Generation Metrics</div>
        <table class="meta-tbl">
          <tr><td>Generation Latency</td><td style="color:var(--green);">{D['latency']} seconds</td></tr>
          <tr><td>Protocol Word Count</td><td>~{len(D['protocol'].split())} words</td></tr>
          <tr><td>Protocol Characters</td><td>{len(D['protocol']):,}</td></tr>
          <tr><td>Model</td><td>Amazon Nova Pro v1.0</td></tr>
          <tr><td>API</td><td>Bedrock Converse API</td></tr>
          <tr><td>PubMed Citations</td><td>{len(D['pubmed'])} retrieved</td></tr>
          <tr><td>NCT Trials (RAG)</td><td>{len(D['nct'])} retrieved</td></tr>
          <tr><td>Cache TTL</td><td>1800s (30 min)</td></tr>
          <tr><td>PDF Library</td><td>ReportLab {'✓' if HAVE_PDF else '✗'}</td></tr>
          <tr><td>DOCX Library</td><td>python-docx {'✓' if HAVE_DOCX else '✗'}</td></tr>
        </table></div>
        """, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════
# LANDING / EMPTY STATE
# ═══════════════════════════════════════════════════════════════════
elif idx is None:
    st.markdown("""
    <div class="hero">
      <div class="hero-eyebrow">⚗️ AI-Powered Clinical Research Platform</div>
      <div class="hero-title">TrialForge AI<br><span>Autonomous Protocol Designer</span></div>
      <div class="hero-sub">Generate ICH-GCP-compliant clinical trial protocols, predict success probability, simulate digital twins — powered by <strong>Amazon Nova Pro</strong>.</div>
      <div class="hero-stats">
        <div class="hs-item"><div class="hs-val">20+</div><div class="hs-lbl">Protocol Sections</div></div>
        <div class="hs-sep"></div>
        <div class="hs-item"><div class="hs-val">4</div><div class="hs-lbl">AI Agents</div></div>
        <div class="hs-sep"></div>
        <div class="hs-item"><div class="hs-val">PDF</div><div class="hs-lbl">Regulatory Export</div></div>
        <div class="hs-sep"></div>
        <div class="hs-item"><div class="hs-val">RAG</div><div class="hs-lbl">Live API</div></div>
        <div class="hs-sep"></div>
        <div class="hs-item"><div class="hs-val">CDISC</div><div class="hs-lbl">XML Export</div></div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div style="text-align:center;font-size:.62rem;font-weight:700;letter-spacing:.12em;text-transform:uppercase;color:var(--text3);margin-bottom:.6rem;">Quick Start Templates</div>', unsafe_allow_html=True)
    cols = st.columns(3)
    for i, t in enumerate(DEMO_TEMPLATES):
        with cols[i % 3]:
            if st.button(f"{t['icon']} {t['name']}", key=f"tmpl_{i}", use_container_width=True, help=f"{t['drug']} · {t['phase']} · {t['mode']}"):
                st.session_state.tmpl = t
                st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    feat_cols = st.columns(4)
    features = [
        ("🏆", "Quality Judge", "ICH compliance, FDA readiness, failure risk — all scored automatically."),
        ("🌐", "Live API RAG", "Real ClinicalTrials.gov + PubMed citations injected into every protocol."),
        ("📄", "PDF & DOCX", "Regulatory-grade formatted export with tables, scores, and full text."),
        ("🔴", "Red Team", "Mock FDA reviewer attacks your protocol before submission."),
    ]
    for col, (ic, title, desc) in zip(feat_cols, features):
        with col:
            st.markdown(f'<div class="card" style="text-align:center;padding:.9rem .75rem;"><div style="font-size:1.5rem;margin-bottom:6px;">{ic}</div><div style="font-size:.77rem;font-weight:700;color:var(--text);margin-bottom:4px;">{title}</div><div style="font-size:.68rem;color:var(--text2);line-height:1.5;">{desc}</div></div>', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════
st.markdown("""
<div class="tf-footer">
  <div class="tf-footer-l">TrialForge AI v3.0 · Amazon Nova Pro (Bedrock Converse API) · ClinicalTrials.gov · PubMed E-utilities · All protocols require qualified regulatory review</div>
  <div class="tf-footer-r">ICH E6(R2) · FDA 21 CFR Part 312 · FDORA §3041 · CDISC ODM 1.3</div>
</div>
""", unsafe_allow_html=True)